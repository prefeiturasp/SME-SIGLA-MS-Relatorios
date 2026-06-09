"""Implementação concreta do relatório de Ata de Escolha.

Baseado no padrão da Ata de Convocação, mas com informações da escola
escolhida.
"""

from __future__ import annotations

import logging
import os
import re
import tempfile
from io import BytesIO
from typing import Any

import requests
from django.conf import settings
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render

from relatorios.services.ata_escolha_service import (
    AtaEscolhaService,
    CargoObrigatorioError,
)
from relatorios.services.base.relatorio_base import RelatorioBase

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
try:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
logger = logging.getLogger(__name__)


class AtaEscolha(RelatorioBase):
    """Classe concreta responsável por gerar o relatório de Ata de Escolha."""

    TEMPLATE_NAME = "relatorios/ata_escolha.html"

    def __init__(self, **kwargs: Any) -> None:
        """Inicializa o service com as dependências necessárias.

        Args:
            self: Instância do objeto.
            **kwargs: Argumentos nomeados variáveis.

        Raises:
            Nenhuma exceção específica documentada.
        """
        super().__init__(**kwargs)
        self.ata_service = AtaEscolhaService(
            candidatos_base_url=settings.CANDIDATOS_API_URL,
            processo_base_url=settings.CONVOCACAO_API_URL,
            agendas_base_url=settings.AGENDAS_API_URL,
            escolhas_base_url=settings.ESCOLHAS_API_URL,
        )

    def _preencher_template(self, cabecalho_capa: Any, dados: Any) -> Any:
        """Executa  preencher template.

        Args:
            self: Instância do objeto.
            cabecalho_capa: Parâmetro cabecalho capa.
            dados: Parâmetro dados.

        Returns:
            Resultado da operação.

        Raises:
            Nenhuma exceção específica documentada.
        """
        pattern = re.compile("\\[\\[(.*?)\\]\\]")

        def replace_func(match: Any) -> Any:
            """Executa replace func.

            Args:
                match: Parâmetro match.

            Returns:
                Resultado da operação.

            Raises:
                Nenhuma exceção específica documentada.
            """
            chave = match.group(1)
            return str(dados.get(chave, f"[[ERRO: {chave} NÃO ENCONTRADO]]"))

        return pattern.sub(replace_func, cabecalho_capa)

    def gerar(
        self,
        processo_uuid: str,
        request: Any,
        formato: str = "html",
        cabecalho: str = "",
        cargo_codigo: str = None,
        **kwargs: Any,
    ) -> Any:  # type: ignore[assignment]
        """Gera o relatório de Ata de Escolha para um único cargo.

        Args:
            self: Instância do objeto.
            processo_uuid: UUID do processo de convocação.
            request: Objeto request do Django.
            formato: Formato do relatório ('html', 'pdf' ou 'xls').
            cabecalho: Texto do cabeçalho do relatório (opcional).
            cargo_codigo: Fixture do teste.
            **kwargs: Argumentos nomeados variáveis.

        Returns:
            Resultado da operação.

        Raises:
            Nenhuma exceção específica documentada.
        """
        try:
            dados_ata = self.ata_service.processar_ata_escolha(
                processo_uuid=str(processo_uuid) if processo_uuid else "",
                cargo_codigo=cargo_codigo or None,
            )
        except CargoObrigatorioError:
            raise
        except Exception as exc:
            logger.error("Falha ao processar ata de escolha: %s", exc)
            raise
        cargos_list = dados_ata.get("cargos", [])
        primeiro_cargo = cargos_list[0] if cargos_list else {}
        datas_preencher_tempalte = {
            "cargo": primeiro_cargo.get("cargo_nome", ""),
            "tipos_vagas": "PRECÁRIAS/DEFINITIVAS",
        }
        cabecalho_capa_ata = self._preencher_template(
            self.context["cabecalho_capa_ata"], datas_preencher_tempalte
        )
        logo_url = (
            request.build_absolute_uri(self.context.get("logo_url", ""))
            if self.context.get("logo_url")
            else ""
        )
        cabecalho_processado = (
            cabecalho.strip() if cabecalho and cabecalho.strip() else None
        )
        context_data = self.context.copy()
        context_data.update(
            {
                "cargos": dados_ata.get("cargos", []),
                "candidatos_sep_cargo": dados_ata.get(
                    "candidatos_sep_cargo", {}
                ),
                "cabecalho_capa_ata": cabecalho_capa_ata,
                "mostrar_capa_ata": False,
                "escolhas_totais_por_tipo": dados_ata.get(
                    "escolhas_totais_por_tipo", {}
                ),
                "logo_url": logo_url,
                "is_pdf": False,
                "intervalos_classificacoes": dados_ata.get(
                    "intervalos_classificacoes", {}
                ),
                "processo_nome": dados_ata.get("processo_nome", ""),
            }
        )
        if cabecalho_processado:
            context_data["cabecalho"] = cabecalho_processado
        if formato == "docx" or formato == "doc":
            filename = f"ata_escolha_{processo_uuid}.docx"
            logger.info("Gerando Word: %s", filename)
            response = self.render_to_docx(
                dados_ata.get("cargos", []), self.context, filename=filename
            )
            return (response, dados_ata)
        elif formato == "pdf":
            context_data["mostrar_capa_ata"] = True
            filename = f"ata_escolha_{processo_uuid}.pdf"
            logger.info("Gerando PDF: %s", filename)
            context_data["is_pdf"] = True
            response = self.render_to_pdf(
                self.TEMPLATE_NAME, context_data, filename=filename
            )
            return (response, dados_ata)
        elif formato in ("xls", "xlsx"):
            filename = f"ata_escolha_{processo_uuid}.xlsx"
            logger.info("Gerando XLS: %s", filename)
            response = self._render_xls(context_data, filename=filename)
            return (response, dados_ata)
        elif formato == "html":
            context_data["mostrar_capa_ata"] = True
            logger.info("Gerando HTML")
            response = render(request, self.TEMPLATE_NAME, context_data)
            return (response, dados_ata)
        else:
            response = JsonResponse(
                dados_ata,
                safe=False,
                json_dumps_params={"indent": 2, "ensure_ascii": False},
            )
        return (response, dados_ata)

    def render_to_docx(
        self,
        cargos_list: Any,
        context: Any,
        filename: Any = "ata_escolha.docx",
    ) -> Any:
        """Gera um arquivo Word (DOCX) mantendo a estrutura hierárquica do.

        Args:
            self: Instância do objeto.
            cargos_list: Lista de cargos com suas sessões e candidatos.
            context: Dicionário com o contexto do template.
            filename: Nome do arquivo Word gerado.

        Returns:
            Resultado da operação.

        Raises:
            ImportError: Se ocorrer erro nesta operação.
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx não está instalado. Instale com: pip install python-docx>=1.1.0"  # noqa: E501
            )
        try:
            doc = Document()
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Cm(50)
            section.page_height = Cm(35)
            RGBColor(52, 73, 94)
            RGBColor(102, 126, 234)
            RGBColor(236, 240, 241)
            if context.get("cabecalho_padrao"):
                cabecalho_texto = self.processar_cabecalho_html(
                    context.get("cabecalho_padrao")
                )
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cabecalho_texto)
                run.font.size = Pt(12)
                run.font.bold = True
                doc.add_paragraph()
            if context.get("cabecalho"):
                cabecalho_texto = self.processar_cabecalho_html(
                    context.get("cabecalho")
                )
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cabecalho_texto)
                run.font.size = Pt(12)
                run.font.bold = True
                doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("ATA DE ESCOLHA")
            run.font.size = Pt(14)
            run.font.bold = True
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("")
            run.font.size = Pt(10)
            run.font.italic = True
            doc.add_paragraph()
            for cargo in cargos_list:
                cargo_nome = cargo.get("cargo_nome", "")
                for sessao in cargo.get("sessoes", []):
                    numero_sessao = sessao.get("numero_sessao", "")
                    horario_formatado = sessao.get("horario_formatado", "")
                    sessao_texto = f"{numero_sessao}ª SESSÃO"
                    if horario_formatado:
                        sessao_texto += f" - Horário: {horario_formatado}"
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(sessao_texto)
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    p_pr = p._element.get_or_add_pPr()
                    existing_shd = p_pr.find(qn("w:shd"))
                    if existing_shd is not None:
                        p_pr.remove(existing_shd)
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:fill"), "34495e")
                    shading_elm.set(qn("w:val"), "clear")
                    p_pr.append(shading_elm)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(f"CARGO: {cargo_nome}")
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    p_pr = p._element.get_or_add_pPr()
                    existing_shd = p_pr.find(qn("w:shd"))
                    if existing_shd is not None:
                        p_pr.remove(existing_shd)
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:fill"), "667eea")
                    shading_elm.set(qn("w:val"), "clear")
                    p_pr.append(shading_elm)
                    headers = [
                        "Class. Geral",
                        "Class. Def.",
                        "Class. NNA",
                        "RF",
                        "RG",
                        "CPF",
                        "Nome",
                        "Código EOL",
                        "DRE",
                        "Tipo Unidade",
                        "Unidade Escolhida",
                        "Tipo Vaga",
                        "Assinatura",
                    ]
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = "Light Grid Accent 1"
                    header_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        cell = header_cells[i]
                        cell.text = header
                        cell.paragraphs[0].alignment = (
                            WD_ALIGN_PARAGRAPH.CENTER
                            if i < 6 or i >= 7
                            else WD_ALIGN_PARAGRAPH.LEFT
                        )
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].runs[0].font.size = Pt(7)
                        tc_pr = cell._element.get_or_add_tcPr()
                        existing_shd = tc_pr.find(qn("w:shd"))
                        if existing_shd is not None:
                            tc_pr.remove(existing_shd)
                        shading_elm = OxmlElement("w:shd")
                        shading_elm.set(qn("w:fill"), "ECF0F1")
                        shading_elm.set(qn("w:val"), "clear")
                        tc_pr.append(shading_elm)
                    for candidato in sessao.get("candidatos", []):
                        row_cells = table.add_row().cells
                        candidato_obj = (
                            candidato.get("candidato", {})
                            if isinstance(candidato.get("candidato"), dict)
                            else {}
                        )
                        row_cells[0].text = (
                            "-"
                            if str(candidato.get("classificacao", "-"))
                            == "None"
                            else str(candidato.get("classificacao", "-"))
                        )
                        row_cells[1].text = (
                            "-"
                            if str(candidato.get("classificacao_pcd", "-"))
                            == "None"
                            else str(candidato.get("classificacao_pcd", "-"))
                        )
                        row_cells[2].text = (
                            "-"
                            if str(candidato.get("classificacao_nna", "-"))
                            == "None"
                            else str(candidato.get("classificacao_nna", "-"))
                        )
                        rf = (
                            candidato.get("rf", "")
                            or (
                                candidato_obj.get("registro_funcional", "")
                                if isinstance(candidato_obj, dict)
                                else ""
                            )
                            or "-"
                        )
                        rg = (
                            candidato.get("rg", "")
                            or (
                                candidato_obj.get("rg", "")
                                if isinstance(candidato_obj, dict)
                                else ""
                            )
                            or "-"
                        )
                        cpf = (
                            candidato.get("cpf", "")
                            or (
                                candidato_obj.get("cpf", "")
                                if isinstance(candidato_obj, dict)
                                else ""
                            )
                            or "-"
                        )
                        row_cells[3].text = str(rf)
                        row_cells[4].text = str(rg)
                        row_cells[5].text = str(cpf)
                        status_especial = candidato.get("status_especial", "")
                        if status_especial:
                            row_cells[6].text = status_especial
                        else:
                            nome = candidato.get("nome", "") or (
                                candidato_obj.get("nome", "N/A")
                                if isinstance(candidato_obj, dict)
                                else "N/A"
                            )
                            row_cells[6].text = nome
                        row_cells[7].text = str(
                            candidato.get("codigo_eol", "-") or "-"
                        )
                        row_cells[8].text = str(
                            candidato.get("dre_codigo", "-") or "-"
                        )
                        row_cells[9].text = str(
                            candidato.get("tipo_unidade", "-") or "-"
                        )
                        row_cells[10].text = str(
                            candidato.get("nome_escola_escolhida", "-") or "-"
                        )
                        row_cells[11].text = str(
                            candidato.get("tipo_vaga", "-") or "-"
                        )
                        row_cells[12].text = str(
                            candidato.get("assinatura", "Não Escolha")
                            or "Não Escolha"
                        )
                        for i, cell in enumerate(row_cells):
                            if i < 6 or (i >= 7 and i < 12):
                                cell.paragraphs[
                                    0
                                ].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                cell.paragraphs[
                                    0
                                ].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            cell.paragraphs[0].runs[0].font.size = Pt(7)
                    doc.add_paragraph()
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            response = HttpResponse(
                buffer.read(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = (
                f'attachment; filename="{filename}"'
            )
            return response
        except Exception as exc:
            logger.error("Erro ao gerar Word: %s", exc, exc_info=True)
            raise

    def _render_xls(
        self, context_data: Any, filename: Any = "ata_escolha.xlsx"
    ) -> Any:
        """Gera um arquivo Excel (XLSX) com a estrutura da Ata de Escolha.

        Args:
            self: Instância do objeto.
            context_data: Parâmetro context data.
            filename: Parâmetro filename.

        Returns:
            Resultado da operação.

        Raises:
            ImportError: Se ocorrer erro nesta operação.
        """
        if not OPENPYXL_AVAILABLE:
            raise ImportError(
                "openpyxl não está instalado. Instale com: pip install openpyxl>=3.1.0"  # noqa: E501
            )
        wb = Workbook()
        ws = wb.active
        ws.title = "Ata de Escolha"
        header_fill = PatternFill(
            start_color="ECF0F1", end_color="ECF0F1", fill_type="solid"
        )
        header_font = Font(bold=True, size=10)
        title_font = Font(bold=True, size=12)
        normal_font = Font(size=9)
        center = Alignment(horizontal="center", vertical="center")
        left = Alignment(horizontal="left", vertical="center")
        border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )
        row_idx = 1
        temp_image_paths = []
        logo_url = (
            (context_data or self.context).get("logo_url")
            if context_data or self.context
            else ""
        )
        if context_data.get("usar_logotipo") and logo_url:
            image_path = None
            try:
                if logo_url.startswith("http://") or logo_url.startswith(
                    "https://"
                ):
                    with tempfile.NamedTemporaryFile(
                        suffix=".png", delete=False
                    ) as tmpf:
                        resp = requests.get(logo_url, timeout=15)
                        resp.raise_for_status()
                        tmpf.write(resp.content)
                        image_path = tmpf.name
                        temp_image_paths.append(image_path)
                elif os.path.exists(logo_url):
                    image_path = logo_url
                if image_path:
                    img = XLImage(image_path)
                    try:
                        img.width = 220
                        img.height = 90
                    except Exception:
                        pass
                    ws.add_image(img, "G1")
                    row_idx = max(row_idx, 8)
            except Exception as exc:
                logger.warning(
                    "Não foi possível inserir o logotipo no XLS: %s", exc
                )
        cabecalho_padrao = self.context.get("cabecalho_padrao", "")
        if cabecalho_padrao:
            ws.merge_cells(f"A{row_idx}:M{row_idx}")
            cell = ws[f"A{row_idx}"]
            cell.value = self.processar_cabecalho_html(cabecalho_padrao)
            cell.font = title_font
            cell.alignment = center
            row_idx += 2
        if self.context.get("cabecalho"):
            ws.merge_cells(f"A{row_idx}:M{row_idx}")
            cell = ws[f"A{row_idx}"]
            cell.value = self.processar_cabecalho_html(
                self.context["cabecalho"]
            )
            cell.font = title_font
            cell.alignment = center
            row_idx += 2
        headers = [
            "Class. Geral",
            "Class. Def.",
            "Class. NNA",
            "RF",
            "RG",
            "CPF",
            "Nome",
            "Código EOL",
            "DRE",
            "Tipo Unidade",
            "Unidade Escolhida",
            "Tipo Vaga",
            "Assinatura",
        ]
        for cargo in context_data.get("cargos", []) or []:
            cargo_nome = cargo.get("cargo_nome", "")
            sessoes = cargo.get("sessoes", []) or []
            for sessao in sessoes:
                numero_sessao = sessao.get("numero_sessao", "")
                horario_formatado = sessao.get("horario_formatado", "")
                sessao_texto = f"{numero_sessao}ª SESSÃO"
                if horario_formatado:
                    sessao_texto += f" - Horário: {horario_formatado}"
                ws.merge_cells(
                    start_row=row_idx,
                    start_column=1,
                    end_row=row_idx,
                    end_column=13,
                )
                c = ws.cell(row=row_idx, column=1)
                c.value = sessao_texto
                c.font = Font(bold=True, size=11)
                c.alignment = left
                c.fill = PatternFill(
                    start_color="34495e", end_color="34495e", fill_type="solid"
                )
                c.font = Font(bold=True, size=11, color="FFFFFF")
                row_idx += 1
                ws.merge_cells(
                    start_row=row_idx,
                    start_column=1,
                    end_row=row_idx,
                    end_column=13,
                )
                c = ws.cell(row=row_idx, column=1)
                c.value = f"CARGO: {cargo_nome}"
                c.font = Font(bold=True, size=11)
                c.alignment = left
                c.fill = PatternFill(
                    start_color="667eea", end_color="667eea", fill_type="solid"
                )
                c.font = Font(bold=True, size=11, color="FFFFFF")
                row_idx += 1
                for col, h in enumerate(headers, start=1):
                    cell = ws.cell(row=row_idx, column=col)
                    cell.value = h
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = (
                        center
                        if col in (1, 2, 3, 4, 5, 6, 8, 9, 10, 12, 13)
                        else left
                    )
                    cell.border = border
                for r, cand in enumerate(
                    sessao.get("candidatos", []), start=row_idx + 1
                ):
                    candidato_obj = (
                        cand.get("candidato", {})
                        if isinstance(cand.get("candidato"), dict)
                        else {}
                    )
                    values = [
                        "-"
                        if str(cand.get("classificacao", "-")) == "None"
                        else str(cand.get("classificacao", "-")),
                        "-"
                        if str(cand.get("classificacao_pcd", "-")) == "None"
                        else str(cand.get("classificacao_pcd", "-")),
                        "-"
                        if str(cand.get("classificacao_nna", "-")) == "None"
                        else str(cand.get("classificacao_nna", "-")),
                    ]
                    rf = (
                        cand.get("rf", "")
                        or (
                            candidato_obj.get("registro_funcional", "")
                            if isinstance(candidato_obj, dict)
                            else ""
                        )
                        or "-"
                    )
                    rg = (
                        cand.get("rg", "")
                        or (
                            candidato_obj.get("rg", "")
                            if isinstance(candidato_obj, dict)
                            else ""
                        )
                        or "-"
                    )
                    cpf = (
                        cand.get("cpf", "")
                        or (
                            candidato_obj.get("cpf", "")
                            if isinstance(candidato_obj, dict)
                            else ""
                        )
                        or "-"
                    )
                    values.extend([str(rf), str(rg), str(cpf)])
                    status_especial = cand.get("status_especial", "")
                    if status_especial:
                        nome = status_especial
                    else:
                        nome = cand.get("nome", "") or (
                            candidato_obj.get("nome", "N/A")
                            if isinstance(candidato_obj, dict)
                            else "N/A"
                        )
                    values.append(nome)
                    values.extend(
                        [
                            str(cand.get("codigo_eol", "-") or "-"),
                            str(cand.get("dre_codigo", "-") or "-"),
                            str(cand.get("tipo_unidade", "-") or "-"),
                            str(cand.get("nome_escola_escolhida", "-") or "-"),
                            str(cand.get("tipo_vaga", "-") or "-"),
                            str(
                                cand.get("assinatura", "Não Escolha")
                                or "Não Escolha"
                            ),
                        ]
                    )
                    for col, val in enumerate(values, start=1):
                        cell = ws.cell(row=r, column=col)
                        cell.value = val
                        cell.font = normal_font
                        cell.alignment = (
                            center
                            if col in (1, 2, 3, 4, 5, 6, 8, 9, 10, 12, 13)
                            else left
                        )
                        cell.border = border
                row_idx = row_idx + 1 + len(sessao.get("candidatos", [])) + 1
        column_widths = {
            "A": 12,
            "B": 12,
            "C": 12,
            "D": 12,
            "E": 15,
            "F": 15,
            "G": 35,
            "H": 12,
            "I": 10,
            "J": 15,
            "K": 40,
            "L": 10,
            "M": 12,
        }
        for col_letter, width in column_widths.items():
            ws.column_dimensions[col_letter].width = width
        texto_final = self.context.get("texto_final")
        if texto_final:
            row_idx += 1
            ws.merge_cells(f"A{row_idx}:D{row_idx}")
            cell = ws[f"A{row_idx}"]
            cell.value = self.processar_cabecalho_html(texto_final)
            cell.font = normal_font
            cell.alignment = Alignment(
                horizontal="left", vertical="top", wrap_text=True
            )
        from io import BytesIO

        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        for p in temp_image_paths:
            try:
                if os.path.exists(p):
                    os.unlink(p)
            except Exception:
                pass
        resp = HttpResponse(
            buf.read(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )  # type: ignore[assignment]
        resp["Content-Disposition"] = f'attachment; filename="{filename}"'  # type: ignore[index]
        return resp

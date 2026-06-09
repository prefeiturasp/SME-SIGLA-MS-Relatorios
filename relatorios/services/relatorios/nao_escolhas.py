"""Módulo services/relatorios/nao_escolhas."""

from __future__ import annotations

import logging
import os
import tempfile
from io import BytesIO
from typing import Any

import requests
from django.conf import settings
from django.http import HttpResponse
from django.shortcuts import render

from relatorios.services.base.relatorio_base import RelatorioBase
from relatorios.services.candidatos_api_service import CandidatosService
from relatorios.services.escolhas_api_service import EscolhasService
from relatorios.services.processos_api_service import ProcessosService
from relatorios.utils import convert_uuids_to_strings

try:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
logger = logging.getLogger(__name__)


class SumulaNaoEscolhas(RelatorioBase):
    """Classe concreta responsável por gerar o relatório de Não Escolhas."""

    TEMPLATE_NAME = "relatorios/nao_escolhas.html"

    def __init__(self, **kwargs: Any) -> None:
        """Inicializa o service com as dependências necessárias.

        Args:
            self: Instância do objeto.
            **kwargs: Argumentos nomeados variáveis.

        Raises:
            Nenhuma exceção específica documentada.
        """
        super().__init__(**kwargs)
        self.escolhas_service = EscolhasService(
            base_url=settings.ESCOLHAS_API_URL
        )
        self.candidatos_service = CandidatosService(
            base_url=settings.CANDIDATOS_API_URL
        )
        self.processos_service = ProcessosService(
            base_url=settings.PROCESSOS_API_URL
        )

    def gerar(
        self,
        processo_uuid: str,
        request: Any,
        formato: str = "html",
        cabecalho: str = "",
        **kwargs: Any,
    ) -> Any:
        """Gera o relatório de Não Escolhas.

        Args:
            self: Instância do objeto.
            processo_uuid: UUID do processo de convocação.
            request: Objeto request do Django.
            formato: Formato do relatório ('html', 'pdf', 'xls' ou 'docx').
            cabecalho: Texto do cabeçalho do relatório (opcional).
            **kwargs: Argumentos nomeados variáveis.

        Returns:
            Resultado da operação.

        Raises:
            Nenhuma exceção específica documentada.
        """
        cargos_map = {}
        try:
            cargos_response = (
                self.processos_service.buscar_cargos_por_processo(
                    processo_uuid=str(processo_uuid) if processo_uuid else ""
                )
            )
            cargos_data = cargos_response.json()
            cargos = cargos_data if isinstance(cargos_data, list) else []
            for cargo in cargos:
                codigo = (
                    cargo.get("cargo_codigo")
                    or cargo.get("codigo_cargo")
                    or ""
                )
                nome = cargo.get("cargo_nome") or cargo.get("nome") or ""
                if codigo and nome:
                    cargos_map[str(codigo)] = nome
                    if isinstance(codigo, int | float):
                        cargos_map[codigo] = nome  # type: ignore[index]
        except Exception as exc:
            logger.warning(
                "Falha ao buscar cargos do processo: %s. Continuando sem mapeamento de cargos.",  # noqa: E501
                exc,
            )
        try:
            candidatos_response = self.candidatos_service.buscar_concurso_candidatos_por_processo(  # noqa: E501
                processo_uuid=str(processo_uuid) if processo_uuid else ""
            )
            candidatos_data = candidatos_response.json()
            candidatos = (
                candidatos_data.get("results", [])
                if isinstance(candidatos_data, dict)
                else candidatos_data
            )
        except Exception as exc:
            logger.error("Falha ao buscar candidatos da API externa: %s", exc)
            raise
        candidatos_map = {}
        for candidato in candidatos:
            candidato_uuid = candidato.get("uuid")
            if candidato_uuid:
                candidatos_map[str(candidato_uuid)] = candidato
        concurso_candidato_uuids = [
            candidato.get("uuid")
            for candidato in candidatos
            if candidato.get("uuid")
        ]
        try:
            escolhas_data = (
                self.escolhas_service.buscar_escolhas_por_candidatos(
                    candidato_uuids=concurso_candidato_uuids,
                    situacao="nao-escolha",
                )
            )
        except Exception as exc:
            logger.error("Falha ao buscar escolhas da API externa: %s", exc)
            raise
        candidatos_com_escolhas = []
        for escolha in escolhas_data:
            candidato_uuid = escolha.get("candidato_uuid")
            if not candidato_uuid:
                continue
            candidato = candidatos_map.get(
                str(candidato_uuid)
            ) or candidatos_map.get(candidato_uuid)
            if not candidato:
                continue
            candidato_obj = (
                candidato.get("candidato", {})
                if isinstance(candidato.get("candidato"), dict)
                else {}
            )
            classificacao_geral = candidato.get("classificacao") or "-"
            classificacao_def = candidato.get("classificacao_pcd") or "-"
            classificacao_nna = candidato.get("classificacao_nna") or "-"
            nome = candidato_obj.get("nome") or "-"
            cpf = candidato_obj.get("cpf") or "-"
            cargo_codigo = candidato.get("codigo_cargo") or ""
            cargo_descricao = candidato.get("descricao_cargo") or ""
            if not cargo_descricao and cargo_codigo:
                cargo_descricao = (
                    cargos_map.get(str(cargo_codigo))
                    or cargos_map.get(cargo_codigo)
                    or ""
                )
            if not cargo_descricao and cargo_codigo:
                cargo_descricao = f"Cargo {cargo_codigo}"
            elif not cargo_descricao:
                cargo_descricao = "Cargo não informado"
            candidatos_com_escolhas.append(
                {
                    "cargo_codigo": cargo_codigo,
                    "cargo_descricao": cargo_descricao,
                    "classificacao_geral": classificacao_geral,
                    "classificacao_def": classificacao_def,
                    "classificacao_nna": classificacao_nna,
                    "nome": nome,
                    "cpf": cpf,
                }
            )
        cargos_list = self._agrupar_por_cargo(candidatos_com_escolhas)
        cargos_list = convert_uuids_to_strings(cargos_list)
        cabecalho_final = self.context.get(
            "cabecalho_padrao", ""
        ) or self.context.get("cabecalho", "")
        logo_url = (
            request.build_absolute_uri(self.context.get("logo_url", ""))
            if self.context.get("logo_url")
            else ""
        )
        self.context.update(
            {"cargos": cargos_list, "is_pdf": False, "logo_url": logo_url}
        )
        if formato == "xls" or formato == "csv":
            filename = f"relatorio_nao_escolhas_{processo_uuid}.xlsx"
            logger.info("Gerando Excel: %s", filename)
            response = self.render_to_xls(
                context=self.context, filename=filename
            )
            return (response, cargos_list)
        elif formato == "docx" or formato == "doc":
            filename = f"relatorio_nao_escolhas_{processo_uuid}.docx"
            logger.info("Gerando Word: %s", filename)
            response = self.render_to_docx(
                cargos_list,
                cabecalho_final,
                self.context["texto_final"],
                filename=filename,
            )
            return (response, cargos_list)
        elif formato == "pdf":
            filename = f"relatorio_nao_escolhas_{processo_uuid}.pdf"
            logger.info("Gerando PDF: %s", filename)
            self.context.update({"cargos": cargos_list})
            response = self.render_to_pdf(
                self.TEMPLATE_NAME, self.context, filename=filename
            )
            return (response, cargos_list)
        else:
            logger.info("Gerando HTML")
            self.context.update({"cargos": cargos_list})
            response = render(request, self.TEMPLATE_NAME, self.context)
            return (response, cargos_list)

    def _agrupar_por_cargo(self, candidatos: list) -> list:
        """Agrupa candidatos por cargo.

        Args:
            self: Instância do objeto.
            candidatos: Lista de candidatos com suas informações.

        Returns:
            Lista com os registros resultantes.

        Raises:
            Nenhuma exceção específica documentada.
        """
        cargos_dict = {}  # type: ignore[var-annotated]
        for candidato in candidatos:
            cargo_codigo = candidato.get("cargo_codigo", "") or ""
            cargo_descricao = candidato.get("cargo_descricao", "") or ""
            if not cargo_descricao:
                if cargo_codigo and cargo_codigo != "-":
                    cargo_descricao = f"Cargo {cargo_codigo}"
                else:
                    cargo_descricao = "Cargo não informado"
            if cargo_codigo not in cargos_dict:
                cargos_dict[cargo_codigo] = {
                    "codigo": cargo_codigo
                    if cargo_codigo and cargo_codigo != "-"
                    else "",
                    "descricao": cargo_descricao,
                    "candidatos": [],
                }
            cargos_dict[cargo_codigo]["candidatos"].append(candidato)
        cargos_list = list(cargos_dict.values())
        cargos_list.sort(key=lambda x: x["descricao"])
        for cargo in cargos_list:
            cargo["candidatos"].sort(
                key=lambda c: c["classificacao_geral"]
                if isinstance(c["classificacao_geral"], int | float)
                else float("inf")
            )
        return cargos_list

    def render_to_xls(
        self,
        context: Any = None,
        filename: Any = "relatorio_nao_escolhas.xlsx",
    ) -> Any:
        """Gera um arquivo Excel (XLSX) mantendo a estrutura hierárquica do.

        Args:
            self: Instância do objeto.
            context: Contexto do relatório.
            filename: Nome do arquivo Excel gerado.

        Returns:
            Resultado da operação.

        Raises:
            ImportError: Se ocorrer erro nesta operação.
        """
        if context is None:
            context = {}
        if not OPENPYXL_AVAILABLE:
            raise ImportError(
                "openpyxl não está instalado. Instale com: pip install openpyxl>=3.1.0"  # noqa: E501
            )
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Relatório de Não Escolhas"
            cargo_fill = PatternFill(
                start_color="667eea", end_color="667eea", fill_type="solid"
            )
            table_header_fill = PatternFill(
                start_color="ECF0F1", end_color="ECF0F1", fill_type="solid"
            )
            cargo_font = Font(bold=True, color="FFFFFF", size=12)
            header_font = Font(bold=True, size=10)
            normal_font = Font(size=10)
            title_font = Font(bold=True, size=14)
            border = Border(
                left=Side(style="thin"),
                right=Side(style="thin"),
                top=Side(style="thin"),
                bottom=Side(style="thin"),
            )
            center_align = Alignment(horizontal="center", vertical="center")
            center_wrap_align = Alignment(
                horizontal="center", vertical="center", wrap_text=True
            )
            left_align = Alignment(horizontal="left", vertical="center")
            row = 1
            temp_image_paths = []
            logo_url = (
                (context or self.context).get("logo_url")
                if context or self.context
                else ""
            )
            if context.get("usar_logotipo") and logo_url:
                image_path = None
                try:
                    if logo_url.startswith("http://") or logo_url.startswith(
                        "https://"
                    ):
                        with tempfile.NamedTemporaryFile(
                            suffix=".png", delete=False
                        ) as tmpf:
                            resp = requests.get(logo_url, timeout=15)
                            resp.raise_for_status()
                            tmpf.write(resp.content)
                            image_path = tmpf.name
                            temp_image_paths.append(image_path)
                    elif os.path.exists(logo_url):
                        image_path = logo_url
                    if image_path:
                        img = XLImage(image_path)
                        try:
                            img.width = 220
                            img.height = 90
                        except Exception:
                            pass
                        ws.add_image(img, "B1")
                        row = max(row, 8)
                except Exception as exc:
                    logger.warning(
                        "Não foi possível inserir o logotipo no XLS (reconvocacao): %s",  # noqa: E501
                        exc,
                    )
            cabecalho_padrao = self.context.get("cabecalho_padrao", "")
            if cabecalho_padrao:
                ws.merge_cells(f"A{row}:D{row}")
                cell = ws[f"A{row}"]
                cell.value = self.processar_cabecalho_html(cabecalho_padrao)
                cell.font = title_font
                cell.alignment = center_wrap_align
                row += 2
            if self.context.get("cabecalho"):
                ws.merge_cells(f"A{row}:D{row}")
                cell = ws[f"A{row}"]
                cell.value = self.processar_cabecalho_html(
                    self.context["cabecalho"]
                )
                cell.font = title_font
                cell.alignment = center_wrap_align
                row += 2
            for cargo in context.get("cargos", []):
                cargo_descricao = cargo.get("descricao", "")
                ws.merge_cells(f"A{row}:D{row}")
                cell = ws[f"A{row}"]
                cell.value = f"Cargo: {cargo_descricao}"
                cell.font = cargo_font
                cell.fill = cargo_fill
                cell.alignment = left_align
                row += 1
                headers = [
                    "Class. Geral",
                    "Class. Def.",
                    "Class. NNA",
                    "Candidato",
                ]
                for col, header in enumerate(headers, start=1):
                    cell = ws.cell(row=row, column=col)
                    cell.value = header
                    cell.fill = table_header_fill
                    cell.font = header_font
                    cell.alignment = center_align
                    cell.border = border
                row += 1
                for candidato in cargo.get("candidatos", []):
                    ws.cell(row=row, column=1).value = candidato.get(
                        "classificacao_geral", "-"
                    )
                    ws.cell(row=row, column=2).value = candidato.get(
                        "classificacao_def", "-"
                    )
                    ws.cell(row=row, column=3).value = candidato.get(
                        "classificacao_nna", "-"
                    )
                    ws.cell(row=row, column=4).value = candidato.get(
                        "nome", "-"
                    )
                    for col in range(1, 5):
                        cell = ws.cell(row=row, column=col)
                        cell.border = border
                        cell.font = normal_font
                        if col in [1, 2, 3]:
                            cell.alignment = center_align
                        else:
                            cell.alignment = left_align
                    row += 1
                row += 1
            column_widths = {"A": 15, "B": 15, "C": 15, "D": 50}
            for col_letter, width in column_widths.items():
                ws.column_dimensions[col_letter].width = width
            texto_final = self.context.get("texto_final")
            if texto_final:
                row += 1
                ws.merge_cells(f"A{row}:D{row}")
                cell = ws[f"A{row}"]
                cell.value = self.processar_cabecalho_html(texto_final)
                cell.font = normal_font
                cell.alignment = Alignment(
                    horizontal="left", vertical="top", wrap_text=True
                )
            buffer = BytesIO()
            wb.save(buffer)
            buffer.seek(0)
            for p in temp_image_paths:
                try:
                    if os.path.exists(p):
                        os.unlink(p)
                except Exception:
                    pass
            response = HttpResponse(
                buffer.read(),
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            response["Content-Disposition"] = (
                f'attachment; filename="{filename}"'
            )
            return response
        except Exception as exc:
            logger.error("Erro ao gerar Excel: %s", exc, exc_info=True)
            raise

    def render_to_docx(
        self,
        cargos_list: Any,
        cabecalho: Any,
        texto_final: Any,
        filename: Any = "relatorio_nao_escolhas.docx",
    ) -> Any:
        """Gera um arquivo Word (DOCX) mantendo a estrutura hierárquica do.

        Args:
            self: Instância do objeto.
            cargos_list: Lista de cargos com seus candidatos (estrutura.
            cabecalho: Texto do cabeçalho do relatório.
            texto_final: Texto final do relatório.
            filename: Nome do arquivo Word gerado.

        Returns:
            Resultado da operação.

        Raises:
            ImportError: Se ocorrer erro nesta operação.
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx não está instalado. Instale com: pip install python-docx>=1.1.0"  # noqa: E501
            )
        try:
            doc = Document()
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            RGBColor(102, 126, 234)
            RGBColor(236, 240, 241)
            for cab in [
                self.context.get("cabecalho_padrao", ""),
                self.context.get("cabecalho", ""),
            ]:
                if cab:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(self.processar_cabecalho_html(cab))
                    run.font.size = Pt(14)
                    run.font.bold = True
                    doc.add_paragraph()
            for cargo in cargos_list:
                cargo_descricao = cargo.get("descricao", "")
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(f"Cargo: {cargo_descricao}")
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                p_pr = p._element.get_or_add_pPr()
                existing_shd = p_pr.find(qn("w:shd"))
                if existing_shd is not None:
                    p_pr.remove(existing_shd)
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:fill"), "667eea")
                shading_elm.set(qn("w:val"), "clear")
                p_pr.append(shading_elm)
                headers = [
                    "Class. Geral",
                    "Class. Def.",
                    "Class. NNA",
                    "Candidato",
                ]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Light Grid Accent 1"
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    cell = header_cells[i]
                    cell.text = header
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
                    tc_pr = cell._element.get_or_add_tcPr()
                    existing_shd = tc_pr.find(qn("w:shd"))
                    if existing_shd is not None:
                        tc_pr.remove(existing_shd)
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:fill"), "ECF0F1")
                    shading_elm.set(qn("w:val"), "clear")
                    tc_pr.append(shading_elm)
                for candidato in cargo.get("candidatos", []):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(
                        candidato.get("classificacao_geral", "-")
                    )
                    row_cells[1].text = str(
                        candidato.get("classificacao_def", "-")
                    )
                    row_cells[2].text = str(
                        candidato.get("classificacao_nna", "-")
                    )
                    row_cells[3].text = str(candidato.get("nome", "-"))
                    for i, cell in enumerate(row_cells):
                        if i in [0, 1, 2]:
                            cell.paragraphs[
                                0
                            ].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            cell.paragraphs[
                                0
                            ].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        cell.paragraphs[0].runs[0].font.size = Pt(10)
                doc.add_paragraph()
            if texto_final:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(self.processar_cabecalho_html(texto_final))
                run.font.size = Pt(10)
                doc.add_paragraph()
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            response = HttpResponse(
                buffer.read(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = (
                f'attachment; filename="{filename}"'
            )
            return response
        except Exception as exc:
            logger.error("Erro ao gerar Word: %s", exc, exc_info=True)
            raise

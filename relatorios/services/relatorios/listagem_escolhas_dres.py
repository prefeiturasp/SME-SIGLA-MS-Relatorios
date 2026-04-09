"""
Implementação concreta do relatório de Listagem de Escolhas por DREs.
"""
import logging
import tempfile
import os
from django.conf import settings
from django.shortcuts import render
from django.http import HttpResponse
from io import BytesIO
import tempfile
import os
import requests
from relatorios.services.base.relatorio_base import RelatorioBase
from relatorios.services.escolhas_api_service import EscolhasService
from relatorios.services.candidatos_api_service import CandidatosService
from relatorios.utils import convert_uuids_to_strings

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.drawing.image import Image as XLImage
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.shared import Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class ListagemEscolhasDres(RelatorioBase):
    """
    Classe concreta responsável por gerar o relatório de Listagem de Escolhas por DREs.
    """
    
    TEMPLATE_NAME = 'relatorios/listagem_escolhas_dres.html'
    
    def __init__(self, **kwargs):
        """Inicializa o service com as dependências necessárias."""
        super().__init__(**kwargs)
        self.escolhas_service = EscolhasService(base_url=settings.ESCOLHAS_API_URL)
        self.candidatos_service = CandidatosService(base_url=settings.CANDIDATOS_API_URL)
    
    def render_to_xls(self, context={}, filename='listagem_escolhas_dres.xlsx'):
        """
        Gera um arquivo Excel (XLSX) com a listagem de escolhas.
        
        Args:
            context: Contexto do relatório
            filename: Nome do arquivo Excel gerado
        
        Returns:
            HttpResponse com o arquivo Excel gerado
        """
        if not OPENPYXL_AVAILABLE:
            raise ImportError(
                "openpyxl não está instalado. Instale com: pip install openpyxl>=3.1.0"
            )
        
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Listagem de Escolhas"
            
            # Estilos
            header_fill = PatternFill(start_color="4a5568", end_color="4a5568", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=10)
            normal_font = Font(size=9)
            title_font = Font(bold=True, size=12)
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
            left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
            center_wrap_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            row = 1
            
            temp_image_paths = []
            # Inserir logotipo no topo, se disponível
            logo_url = (context or self.context).get('logo_url') if context or self.context else ''
            if context.get('usar_logotipo') and logo_url:
                image_path = None
                try:
                    if logo_url.startswith('http://') or logo_url.startswith('https://'):
                        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmpf:
                            resp = requests.get(logo_url, timeout=15)
                            resp.raise_for_status()
                            tmpf.write(resp.content)
                            image_path = tmpf.name
                            temp_image_paths.append(image_path)
                    elif os.path.exists(logo_url):
                        image_path = logo_url
                    if image_path:
                        img = XLImage(image_path)
                        # opcional: ajustar tamanho
                        try:
                            # Reduz o tamanho da imagem
                            img.width = 220
                            img.height = 90
                        except Exception:
                            pass
                        # Aproxima o alinhamento central ancorando em uma coluna intermediária
                        # Como a planilha usa 4 colunas (A:D), ancorar em B1 fica visualmente centralizado
                        ws.add_image(img, 'B1')
                        # Avança algumas linhas para não sobrepor conteúdo
                        row = max(row, 8)
                except Exception as exc:
                    logger.warning('Não foi possível inserir o logotipo no XLS (listagem_escolhas_dres): %s', exc)

            # Cabeçalho
            cabecalho = self.context['cabecalho_padrao'] if self.context['usar_cabecalho_padrao'] else self.context['cabecalho']
            if cabecalho:
                cabecalho_texto = self.processar_cabecalho_html(cabecalho)
                ws.merge_cells(f'A{row}:O{row}')
                cell = ws[f'A{row}']
                cell.value = cabecalho_texto
                cell.font = title_font
                cell.alignment = center_wrap_align
                row += 2
            
            # Título
            ws.merge_cells(f'A{row}:O{row}')
            cell = ws[f'A{row}']
            cell.value = "Listagem de Escolhas por DREs"
            cell.font = title_font
            cell.alignment = center_align
            row += 2
            
            # Cabeçalhos da tabela
            headers = [
                'Cargo', 'Class', 'Def', 'NNA', 'RF', 'RG', 'CPF',
                'Inscrição', 'Nome', 'Telefone', 'DRE', 'Código EOL',
                'Tipo da unidade', 'Unidade', 'Tipo da vaga'
            ]
            
            for col, header in enumerate(headers, start=1):
                cell = ws.cell(row=row, column=col)
                cell.value = header
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = center_align if col in [2, 3, 4, 12, 15] else left_align
                cell.border = border
            
            # Salvar linha do cabeçalho para freeze_panes
            header_row = row
            row += 1
            
            # Dados das escolhas
            for item in context.get('escolhas', []):
                ws.cell(row=row, column=1).value = item.get('cargo', '-')
                ws.cell(row=row, column=2).value = item.get('classificacao', '-')
                ws.cell(row=row, column=3).value = item.get('classificacao_deficiente', '-')
                ws.cell(row=row, column=4).value = item.get('classificacao_nna', '-')
                ws.cell(row=row, column=5).value = item.get('rf', '-')
                ws.cell(row=row, column=6).value = item.get('rg', '-')
                ws.cell(row=row, column=7).value = item.get('cpf', '-')
                ws.cell(row=row, column=8).value = item.get('inscricao', '-')
                ws.cell(row=row, column=9).value = item.get('nome', '-')
                ws.cell(row=row, column=10).value = item.get('telefone', '-')
                ws.cell(row=row, column=11).value = item.get('dre', '-')
                ws.cell(row=row, column=12).value = item.get('codigo_eol', '-')
                ws.cell(row=row, column=13).value = item.get('tipo_ue', '-')
                ws.cell(row=row, column=14).value = item.get('unidade', '-')
                
                # Tipo da vaga com estilo especial
                tipo_vaga = item.get('tipo_vaga', '-')
                cell_vaga = ws.cell(row=row, column=15)
                cell_vaga.value = tipo_vaga
                if tipo_vaga == 'D':
                    cell_vaga.font = Font(bold=True, size=9, color="2d5016")
                elif tipo_vaga == 'P':
                    cell_vaga.font = Font(bold=True, size=9, color="d97706")
                
                # Aplicar estilos a todas as células da linha
                for col in range(1, 16):
                    cell = ws.cell(row=row, column=col)
                    cell.border = border
                    cell.font = normal_font
                    if col in [2, 3, 4, 12, 15]:  # Colunas centralizadas
                        cell.alignment = center_align
                    else:
                        cell.alignment = left_align
                
                row += 1
            
            # Rodapé com total
            ws.merge_cells(f'A{row}:O{row}')
            cell = ws[f'A{row}']
            cell.value = f"Total de escolhas: {len(context.get('escolhas', []))}"
            cell.font = Font(bold=True, size=9)
            cell.alignment = Alignment(horizontal='right', vertical='center')
            row += 1
            
            # Ajustar largura das colunas
            column_widths = {
                'A': 25,   # Cargo
                'B': 8,    # Class
                'C': 8,    # Def
                'D': 8,    # NNA
                'E': 12,   # RF
                'F': 15,   # RG
                'G': 15,   # CPF
                'H': 12,   # Inscrição
                'I': 30,   # Nome
                'J': 15,   # Telefone
                'K': 35,   # DRE
                'L': 12,   # Código EOL
                'M': 15,   # Tipo da unidade
                'N': 40,   # Unidade
                'O': 10,   # Tipo da vaga
            }
            
            for col_letter, width in column_widths.items():
                ws.column_dimensions[col_letter].width = width
            
            # Congelar primeira linha de dados (cabeçalhos da tabela)
            ws.freeze_panes = f'A{header_row + 1}'

            if context.get('texto_final'):
                row += 1
                ws.merge_cells(f'A{row}:O{row}')
                cell = ws[f'A{row}']
                cell.value = self.processar_cabecalho_html(context.get('texto_final'))
                cell.font = normal_font
                cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
            
            buffer = BytesIO()
            wb.save(buffer)
            buffer.seek(0)
            
            # Limpar temporários de imagem
            for p in temp_image_paths:
                try:
                    if os.path.exists(p):
                        os.unlink(p)
                except Exception:
                    pass

            response = HttpResponse(
                buffer.read(),
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
            
        except Exception as exc:
            logger.error('Erro ao gerar Excel: %s', exc, exc_info=True)
            raise
    
    def render_to_docx(self, escolhas_list, cabecalho, texto_final, filename='listagem_escolhas_dres.docx'):
        """
        Gera um arquivo Word (DOCX) com a listagem de escolhas.
        
        Args:
            escolhas_list: Lista de escolhas com dados dos candidatos
            cabecalho: Texto do cabeçalho do relatório
            texto_final: Texto final do relatório
            filename: Nome do arquivo Word gerado
        
        Returns:
            HttpResponse com o arquivo Word gerado
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx não está instalado. Instale com: pip install python-docx>=1.1.0"
            )
        
        try:
            doc = Document()
            
            # Configurar margens da página
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Cm(42)
            section.page_height = Cm(29.7)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
            # Cores (em RGB)
            table_header_color = RGBColor(74, 85, 104)  # #4a5568
            
            # Cabeçalho
            cabecalho = self.context['cabecalho_padrao'] if self.context['usar_cabecalho_padrao'] else self.context['cabecalho']
            if cabecalho:
                cabecalho_texto = self.processar_cabecalho_html(cabecalho)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cabecalho_texto)
                run.font.size = Pt(12)
                run.font.bold = True
                doc.add_paragraph()
            
            # Título
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("Listagem de Escolhas por DREs")
            run.font.size = Pt(14)
            run.font.bold = True
            doc.add_paragraph()
            
            # Criar tabela
            headers = [
                'Cargo', 'Class', 'Def', 'NNA', 'RF', 'RG', 'CPF',
                'Inscrição', 'Nome', 'Telefone', 'DRE', 'Código EOL',
                'Tipo da unidade', 'Unidade', 'Tipo da vaga'
            ]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Light Grid Accent 1'
            
            # Cabeçalho da tabela
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                cell = header_cells[i]
                cell.text = header
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [1, 2, 3, 11, 14] else WD_ALIGN_PARAGRAPH.LEFT
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(7)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                tc_pr = cell._element.get_or_add_tcPr()
                existing_shd = tc_pr.find(qn('w:shd'))
                if existing_shd is not None:
                    tc_pr.remove(existing_shd)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '4a5568')
                shading_elm.set(qn('w:val'), 'clear')
                tc_pr.append(shading_elm)
            
            # Dados das escolhas
            for item in escolhas_list:
                row_cells = table.add_row().cells
                
                row_cells[0].text = str(item.get('cargo', '-'))
                row_cells[1].text = str(item.get('classificacao', '-'))
                row_cells[2].text = str(item.get('classificacao_deficiente', '-'))
                row_cells[3].text = str(item.get('classificacao_nna', '-'))
                row_cells[4].text = str(item.get('rf', '-'))
                row_cells[5].text = str(item.get('rg', '-'))
                row_cells[6].text = str(item.get('cpf', '-'))
                row_cells[7].text = str(item.get('inscricao', '-'))
                row_cells[8].text = str(item.get('nome', '-'))
                row_cells[9].text = str(item.get('telefone', '-'))
                row_cells[10].text = str(item.get('dre', '-'))
                row_cells[11].text = str(item.get('codigo_eol', '-'))
                row_cells[12].text = str(item.get('tipo_ue', '-'))
                row_cells[13].text = str(item.get('unidade', '-'))
                
                # Tipo da vaga com estilo especial
                tipo_vaga = item.get('tipo_vaga', '-')
                row_cells[14].text = str(tipo_vaga)
                if tipo_vaga == 'D':
                    row_cells[14].paragraphs[0].runs[0].font.bold = True
                    row_cells[14].paragraphs[0].runs[0].font.color.rgb = RGBColor(45, 80, 22)  # #2d5016
                elif tipo_vaga == 'P':
                    row_cells[14].paragraphs[0].runs[0].font.bold = True
                    row_cells[14].paragraphs[0].runs[0].font.color.rgb = RGBColor(217, 119, 6)  # #d97706
                
                # Alinhamento
                for i, cell in enumerate(row_cells):
                    if i in [1, 2, 3, 11, 14]:  # Colunas centralizadas
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    cell.paragraphs[0].runs[0].font.size = Pt(7)
            
            # Rodapé com total
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(f"Total de escolhas: {len(escolhas_list)}")
            run.font.size = Pt(9)
            run.font.bold = True

            if texto_final:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(self.processar_cabecalho_html(texto_final))
                run.font.size = Pt(10)
                doc.add_paragraph()
            
            # Salvar em buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            response = HttpResponse(
                buffer.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
            
        except Exception as exc:
            logger.error('Erro ao gerar Word: %s', exc, exc_info=True)
            raise
    
    def gerar(self, processo_uuid: str, request, formato: str = 'html', cabecalho: str = '', **kwargs):
        """
        Gera o relatório de Listagem de Escolhas por DREs.
        
        Args:
            processo_uuid: UUID do processo de convocação
            request: Objeto request do Django
            formato: Formato do relatório ('html', 'pdf' ou 'xls')
            cabecalho: Texto do cabeçalho do relatório (opcional)
        
        Returns:
            Tupla (HttpResponse, dados) onde:
            - HttpResponse: resposta com o relatório gerado (HTML, PDF ou XLS)
            - dados: estrutura de dados do relatório para salvar no banco
        """
        # 1. Buscar candidatos por processo_uuid
        try:
            logger.info('Buscando candidatos para processo_uuid=%s', processo_uuid)
            candidatos_response = self.candidatos_service.buscar_concurso_candidatos_por_processo(
                processo_uuid=str(processo_uuid) if processo_uuid else '',
            )
            candidatos_data = candidatos_response.json()
            candidatos = candidatos_data.get('results', []) if isinstance(candidatos_data, dict) else candidatos_data
            
            logger.info('Total de candidatos encontrados: %d', len(candidatos))
        except Exception as exc:
            logger.error('Falha ao buscar candidatos da API externa: %s', exc)
            raise
        
        # Criar mapa de candidatos por uuid para busca rápida
        candidatos_map = {}
        for candidato in candidatos:
            candidato_uuid = candidato.get('uuid')
            if candidato_uuid:
                candidatos_map[str(candidato_uuid)] = candidato
        
        # Extrair UUIDs dos candidatos para buscar escolhas
        candidato_uuids = [candidato.get('uuid') for candidato in candidatos if candidato.get('uuid')]
        
        # 2. Buscar escolhas dos candidatos filtrando por situação='escolha'
        try:
            logger.info('Buscando escolhas para %d candidatos com situação=escolha', len(candidato_uuids))
            escolhas_data = self.escolhas_service.buscar_escolhas_por_candidatos(
                candidato_uuids=candidato_uuids,
                situacao='escolha'
            )
            logger.info('Total de escolhas encontradas: %d', len(escolhas_data))
        except Exception as exc:
            logger.error('Falha ao buscar escolhas da API externa: %s', exc)
            raise
        
        # Processar escolhas e associar com dados dos candidatos
        escolhas_com_candidatos = []
        for escolha in escolhas_data:
            candidato_uuid = escolha.get('candidato_uuid')
            if not candidato_uuid:
                continue

            candidato = candidatos_map.get(str(candidato_uuid)) or candidatos_map.get(candidato_uuid)
            if not candidato:
                logger.warning('Candidato UUID %s não encontrado no mapa', candidato_uuid)
                continue
            
            # Extrair dados do candidato
            candidato_obj = candidato.get('candidato', {}) if isinstance(candidato.get('candidato'), dict) else {}
            
            # Obter dados do candidato
            nome = candidato_obj.get('nome') or '-'
            classificacao = candidato.get('classificacao') or '-'
            classificacao_deficiente = candidato.get('classificacao_pcd') or '-'
            classificacao_nna = candidato.get('classificacao_nna') or '-'
            cpf = candidato_obj.get('cpf') or '-'
            rg = candidato_obj.get('rg') or '-'
            telefone = candidato_obj.get('telefone') or '-'
            registro_funcional = candidato_obj.get('registro_funcional') or '-'
            inscricao = candidato.get('inscricao') or candidato.get('numero_inscricao') or '-'
            
            # Extrair dados da vaga_escola (escola e DRE)
            vaga_escola = escolha.get('vaga_escola', {})
            escola = vaga_escola.get('escola', {}) if isinstance(vaga_escola, dict) else {}
            dre = escola.get('dre', {}) if isinstance(escola, dict) else {}
            
            # Dados da vaga e escola
            cargo_descricao = vaga_escola.get('cargo_descricao', '-') if isinstance(vaga_escola, dict) else '-'
            nome_oficial = escola.get('nome_oficial', '-') if isinstance(escola, dict) else '-'
            dre_nome = dre.get('nome', '-') if isinstance(dre, dict) else '-'
            codigo_eol = escola.get('codigo_eol', '-') if isinstance(escola, dict) else '-'
            tipo_ue = escola.get('tipo_ue', '-') if isinstance(escola, dict) else '-'
            
            # Converter tipo_vaga: 'precaria' = 'P', 'definitiva' = 'D'
            tipo_vaga_raw = escolha.get('tipo_vaga', '')
            if tipo_vaga_raw == 'precaria':
                tipo_vaga = 'P'
            elif tipo_vaga_raw == 'definitiva':
                tipo_vaga = 'D'
            else:
                tipo_vaga = '-'

            escolhas_com_candidatos.append({
                'cargo': cargo_descricao,
                'classificacao': classificacao if classificacao != '-' else '',
                'classificacao_deficiente': classificacao_deficiente if classificacao_deficiente != '-' else '',
                'classificacao_nna': classificacao_nna if classificacao_nna != '-' else '',
                'rf': registro_funcional if registro_funcional != '-' else '',
                'rg': rg if rg != '-' else '',
                'cpf': cpf if cpf != '-' else '',
                'inscricao': inscricao if inscricao != '-' else '',
                'nome': nome,
                'telefone': telefone if telefone != '-' else '',
                'dre': dre_nome,
                'codigo_eol': codigo_eol,
                'tipo_ue': tipo_ue,
                'unidade': nome_oficial,
                'tipo_vaga': tipo_vaga,
                # Manter dados originais para compatibilidade
                'escolha': escolha,
                'candidato': candidato,
            })
        
        logger.info('Total de escolhas processadas: %d', len(escolhas_com_candidatos))
        
        # Organizar dados para o template
        # Agrupar por cargo para organização no relatório
        cargos_dict = {}
        for item in escolhas_com_candidatos:
            cargo = item.get('cargo', 'Cargo não informado')
            if cargo not in cargos_dict:
                cargos_dict[cargo] = []
            cargos_dict[cargo].append(item)
        
        # Ordenar itens por cargo, classificação, DRE e unidade
        cargos_list = []
        for cargo, items in sorted(cargos_dict.items()):
            # Ordenar por classificação, DRE e unidade dentro de cada cargo
            # (convertendo classificação para int quando possível)
            items_ordenados = sorted(items, key=lambda x: (
                int(x.get('classificacao', 0)) if x.get('classificacao') and str(x.get('classificacao')).isdigit() else float('inf'),  # Classificação
                x.get('dre', ''),  # DRE
                x.get('unidade', '')  # Unidade
            ))
            cargos_list.append({
                'descricao': cargo,
                'escolhas': items_ordenados
            })
        
        # Obter cabeçalho: prioriza o enviado no request; se vier vazio, usa o padrão do settings
        if cabecalho is not None:
            self.context['cabecalho'] = cabecalho
        cabecalho_final = self.context['cabecalho_padrao'] if self.context['usar_cabecalho_padrao'] else self.context['cabecalho']
        self.context['cabecalho'] = cabecalho_final
        logo_url = request.build_absolute_uri(self.context.get('logo_url', '')) if self.context.get('logo_url') else ''
        
        # Preparar dados para salvar no banco
        dados = {
            'processo_uuid': processo_uuid,
            'total_escolhas': len(escolhas_com_candidatos),
            'escolhas': escolhas_com_candidatos
        }
        
        # Converter todos os UUIDs para strings para garantir serialização JSON
        dados = convert_uuids_to_strings(dados)
        
        # Criar lista ordenada para exportação XLS (lista plana ordenada por cargo, classificação, DRE e unidade)
        escolhas_ordenadas_export = []
        for cargo_item in cargos_list:
            escolhas_ordenadas_export.extend(cargo_item['escolhas'])
        
        # Preparar contexto comum para todos os formatos
        self.context.update({
            'cargos': cargos_list,
            'total_escolhas': len(escolhas_com_candidatos),
            'logo_url': logo_url,
            'is_pdf': False,
            'escolhas': escolhas_ordenadas_export,
        })
        
        if formato == 'xls' or formato == 'xlsx' or formato == 'csv':
            filename = f'listagem_escolhas_dres_{processo_uuid}.xlsx'
            logger.info('Gerando Excel: %s', filename)
            response = self.render_to_xls(context=self.context, filename=filename)
            return response, dados
        elif formato == 'docx' or formato == 'doc':            
            filename = f'listagem_escolhas_dres_{processo_uuid}.docx'
            logger.info('Gerando Word: %s', filename)
            response = self.render_to_docx(escolhas_ordenadas_export, cabecalho_final, self.context.get('texto_final', ''), filename=filename)
            return response, dados
        elif formato == 'pdf':
            filename = f'listagem_escolhas_dres_{processo_uuid}.pdf'
            logger.info('Gerando PDF: %s', filename)
            self.context.update({
                'is_pdf': True,
            })
            response = self.render_to_pdf(
                self.TEMPLATE_NAME,
                self.context,
                filename=filename
            )
            return response, dados
        elif formato == 'html':
            # Gerar HTML
            response = render(request, self.TEMPLATE_NAME, self.context)
            return response, dados
        else:
            # Para outros formatos, retornar JSON por enquanto
            from django.http import JsonResponse
            response = JsonResponse(dados, safe=False)
            return response, dados

"""Implementação concreta do relatório de Lauda de Convocação."""

from __future__ import annotations

import logging
import os
import tempfile
from io import BytesIO
from typing import Any

import requests
from django.conf import settings
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render

from relatorios.services.base.relatorio_base import RelatorioBase
from relatorios.services.lauda_convocacao_service import LaudaConvocacaoService

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
try:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
logger = logging.getLogger(__name__)


class LaudaConvocacao(RelatorioBase):
    """Classe concreta responsável por gerar o relatório de Lauda de."""

    TEMPLATE_NAME = "relatorios/lauda_convocacao.html"

    def __init__(self, **kwargs: Any) -> None:
        """Inicializa a instância com os parâmetros informados.

        Args:
            self: Instância do objeto.
            **kwargs: Argumentos nomeados variáveis.
        """
        super().__init__(**kwargs)
        self.lauda_service = LaudaConvocacaoService(
            candidatos_base_url=settings.CANDIDATOS_API_URL,
            processo_base_url=settings.CONVOCACAO_API_URL,
            agendas_base_url=settings.AGENDAS_API_URL,
            escolhas_base_url=settings.ESCOLHAS_API_URL,
        )

    def gerar(
        self,
        processo_uuid: str,
        request: Any,
        formato: str = "html",
        cabecalho: str = "",
        **kwargs: Any,
    ) -> Any:
        """Gera o relatório de Lauda de Convocação.

        Args:
            self: Instância do objeto.
            processo_uuid: UUID do processo de convocação.
            request: Requisição HTTP recebida.
            formato: Formato utilizado na operação.
            cabecalho: Cabecalho utilizado na operação.
            **kwargs: Argumentos nomeados variáveis.

        Returns:
            Valor calculado conforme a regra aplicada.
        """
        try:
            dados_lauda = self.lauda_service.processar_lauda_convocacao(
                processo_uuid=str(processo_uuid) if processo_uuid else ""
            )
        except Exception as exc:
            logger.error("Falha ao processar lauda de convocação: %s", exc)
            raise
        if cabecalho:
            self.context["cabecalho"] = cabecalho
        logo_url = (
            request.build_absolute_uri(self.context.get("logo_url", ""))
            if self.context.get("logo_url")
            else ""
        )
        self.context["is_pdf"] = False
        self.context["logo_url"] = logo_url
        if formato == "docx" or formato == "doc":
            filename = f"lauda_convocacao_{processo_uuid}.docx"
            logger.info("Gerando Word: %s", filename)
            response = self.render_to_docx(
                dados_lauda.get("cargos", []),
                self.context,
                self.context["texto_final"],
                filename=filename,
            )
            return (response, dados_lauda)
        elif formato == "pdf":
            filename = f"lauda_convocacao_{processo_uuid}.pdf"
            logger.info("Gerando PDF: %s", filename)
            self.context.update(
                {"is_pdf": True, "cargos": dados_lauda.get("cargos", [])}
            )
            response = self.render_to_pdf(
                self.TEMPLATE_NAME, self.context, filename=filename
            )
            return (response, dados_lauda)
        elif formato in ("xls", "xlsx"):
            filename = f"lauda_convocacao_{processo_uuid}.xlsx"
            logger.info("Gerando XLS: %s", filename)
            response = self._render_xls(
                dados_lauda.get("cargos", []),
                context=self.context,
                filename=filename,
            )
            return (response, dados_lauda)
        elif formato == "html":
            logger.info("Gerando HTML")
            self.context.update({"cargos": dados_lauda.get("cargos", [])})
            response = render(request, self.TEMPLATE_NAME, self.context)
            return (response, dados_lauda)
        else:
            response = JsonResponse(dados_lauda, safe=False)
        return (response, dados_lauda)

    def render_to_docx(
        self,
        cargos_list: Any,
        context: Any,
        texto_final: Any,
        filename: Any = "lauda_convocacao.docx",
    ) -> Any:
        """Gera um arquivo Word (DOCX) mantendo a estrutura hierárquica do.

        Args:
            self: Instância do objeto.
            cargos_list: Lista de cargos do processo.
            context: Contexto de serialização ou renderização.
            texto_final: Texto final utilizado na operação.
            filename: Filename utilizado na operação.

        Returns:
            Valor calculado conforme a regra aplicada.

        Raises:
            ImportError: Se ocorrer erro nesta operação.
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx não está instalado. Instale com: pip install python-docx>=1.1.0"  # noqa: E501
            )
        try:
            doc = Document()
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            RGBColor(52, 73, 94)
            RGBColor(102, 126, 234)
            RGBColor(236, 240, 241)
            if context.get("cabecalho_padrao"):
                cabecalho_texto = self.processar_cabecalho_html(
                    context.get("cabecalho_padrao")
                )
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cabecalho_texto)
                run.font.size = Pt(14)
                run.font.bold = True
                doc.add_paragraph()
            if context.get("cabecalho"):
                cabecalho_texto = self.processar_cabecalho_html(
                    context.get("cabecalho")
                )
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cabecalho_texto)
                run.font.size = Pt(14)
                run.font.bold = True
                doc.add_paragraph()
            for cargo in cargos_list:
                cargo_nome = cargo.get("cargo_nome", "")
                for sessao in cargo.get("sessoes", []):
                    numero_sessao = sessao.get("numero_sessao", "")
                    horario_formatado = sessao.get("horario_formatado", "")
                    sessao_texto = f"{numero_sessao}ª SESSÃO"
                    if horario_formatado:
                        sessao_texto += f" - Horário: {horario_formatado}"
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(sessao_texto)
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    p_pr = p._element.get_or_add_pPr()
                    existing_shd = p_pr.find(qn("w:shd"))
                    if existing_shd is not None:
                        p_pr.remove(existing_shd)
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:fill"), "34495e")
                    shading_elm.set(qn("w:val"), "clear")
                    p_pr.append(shading_elm)
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(f"CARGO: {cargo_nome}")
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    p_pr = p._element.get_or_add_pPr()
                    existing_shd = p_pr.find(qn("w:shd"))
                    if existing_shd is not None:
                        p_pr.remove(existing_shd)
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:fill"), "667eea")
                    shading_elm.set(qn("w:val"), "clear")
                    p_pr.append(shading_elm)
                    headers = [
                        "Ordem de Escolha",
                        "Inscrição",
                        "Nome",
                        "Class. Geral",
                        "Class. PcD",
                        "Class. NNA",
                    ]
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = "Light Grid Accent 1"
                    header_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        cell = header_cells[i]
                        cell.text = header
                        cell.paragraphs[0].alignment = (
                            WD_ALIGN_PARAGRAPH.CENTER
                            if i in [0, 1, 3, 4, 5]
                            else WD_ALIGN_PARAGRAPH.LEFT
                        )
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].runs[0].font.size = Pt(10)
                        tc_pr = cell._element.get_or_add_tcPr()
                        existing_shd = tc_pr.find(qn("w:shd"))
                        if existing_shd is not None:
                            tc_pr.remove(existing_shd)
                        shading_elm = OxmlElement("w:shd")
                        shading_elm.set(qn("w:fill"), "ECF0F1")
                        shading_elm.set(qn("w:val"), "clear")
                        tc_pr.append(shading_elm)
                    for candidato in sessao.get("candidatos", []):
                        row_cells = table.add_row().cells
                        ordem_escolha = candidato.get("ordem_escolha", "")
                        row_cells[0].text = (
                            f"{ordem_escolha}º" if ordem_escolha else "-"
                        )
                        codigo_inscricao = candidato.get(
                            "codigo_inscricao"
                        ) or candidato.get("uuid", "-")
                        row_cells[1].text = str(codigo_inscricao)
                        status_especial = candidato.get("status_especial", "")
                        if status_especial:
                            row_cells[2].text = status_especial
                        else:
                            candidato_obj = candidato.get("candidato", {})
                            nome = (
                                candidato_obj.get("nome", "N/A")
                                if isinstance(candidato_obj, dict)
                                else "N/A"
                            )
                            categoria_efetiva = candidato.get(
                                "categoria_efetiva", ""
                            )
                            if categoria_efetiva == "NNA" and candidato.get(
                                "classificacao_nna"
                            ):
                                nome += " (NNA)"
                            elif categoria_efetiva == "PCD" and candidato.get(
                                "classificacao_pcd"
                            ):
                                nome += " (PCD)"
                            row_cells[2].text = nome
                        row_cells[3].text = (
                            "-"
                            if str(candidato.get("classificacao", "-"))
                            == "None"
                            else str(candidato.get("classificacao", "-"))
                        )
                        row_cells[4].text = (
                            "-"
                            if str(candidato.get("classificacao_pcd", "-"))
                            == "None"
                            else str(candidato.get("classificacao_pcd", "-"))
                        )
                        row_cells[5].text = (
                            "-"
                            if str(candidato.get("classificacao_nna", "-"))
                            == "None"
                            else str(candidato.get("classificacao_nna", "-"))
                        )
                        for i, cell in enumerate(row_cells):
                            if i in [0, 1, 3, 4, 5]:
                                cell.paragraphs[
                                    0
                                ].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                cell.paragraphs[
                                    0
                                ].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            cell.paragraphs[0].runs[0].font.size = Pt(10)
                    doc.add_paragraph()
            if texto_final:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(self.processar_cabecalho_html(texto_final))
                run.font.size = Pt(10)
                doc.add_paragraph()
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            response = HttpResponse(
                buffer.read(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = (
                f'attachment; filename="{filename}"'
            )
            return response
        except Exception as exc:
            logger.error("Erro ao gerar Word: %s", exc, exc_info=True)
            raise

    def _render_xls(
        self,
        cargos_list: Any,
        context: Any = None,
        filename: Any = "lauda_convocacao.xlsx",
    ) -> Any:
        """Gera planilha XLSX da lauda de convocação.

        Args:
            self: Instância do objeto.
            cargos_list: Lista de cargos do processo.
            context: Contexto de serialização ou renderização.
            filename: Filename utilizado na operação.

        Returns:
            Valor calculado conforme a regra aplicada.

        Raises:
            ImportError: Se ocorrer erro nesta operação.
        """
        if context is None:
            context = {}
        if not OPENPYXL_AVAILABLE:
            raise ImportError(
                "openpyxl não está instalado. Instale com: pip install openpyxl>=3.1.0"  # noqa: E501
            )
        wb = Workbook()
        ws = wb.active
        ws.title = "Lauda de Convocação"
        header_fill = PatternFill(
            start_color="ECF0F1", end_color="ECF0F1", fill_type="solid"
        )
        header_font = Font(bold=True, size=11)
        title_font = Font(bold=True, size=14)
        label_font = Font(bold=True, size=12)
        normal_font = Font(size=10)
        center = Alignment(horizontal="center", vertical="center")
        left = Alignment(horizontal="left", vertical="center")
        border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )
        temp_image_paths = []
        row_idx = 1
        logo_url = (
            (context or self.context).get("logo_url")
            if context or self.context
            else ""
        )
        if context.get("usar_logotipo") and logo_url:
            image_path = None
            try:
                if logo_url.startswith("http://") or logo_url.startswith(
                    "https://"
                ):
                    with tempfile.NamedTemporaryFile(
                        suffix=".png", delete=False
                    ) as tmpf:
                        resp = requests.get(logo_url, timeout=15)
                        resp.raise_for_status()
                        tmpf.write(resp.content)
                        image_path = tmpf.name
                        temp_image_paths.append(image_path)
                elif os.path.exists(logo_url):
                    image_path = logo_url
                if image_path:
                    img = XLImage(image_path)
                    try:
                        img.width = 220
                        img.height = 90
                    except Exception:
                        pass
                    ws.add_image(img, "C1")
                    row_idx = max(row_idx, 8)
            except Exception as exc:
                logger.warning(
                    "Não foi possível inserir o logotipo no XLS (lauda_convocacao): %s",  # noqa: E501
                    exc,
                )
        ws.merge_cells(
            start_row=row_idx, start_column=1, end_row=row_idx, end_column=6
        )
        c = ws.cell(row=row_idx, column=1)
        c.value = "Lauda de Convocação"
        c.font = title_font
        c.alignment = center
        row_idx += 1
        cabecalho_padrao = self.context.get("cabecalho_padrao", "")
        if cabecalho_padrao:
            ws.merge_cells(
                start_row=row_idx,
                start_column=1,
                end_row=row_idx,
                end_column=6,
            )
            c = ws.cell(row=row_idx, column=1)
            c.value = self.processar_cabecalho_html(cabecalho_padrao)
            c.font = label_font
            c.alignment = left
            row_idx += 1
        if self.context.get("cabecalho"):
            ws.merge_cells(
                start_row=row_idx,
                start_column=1,
                end_row=row_idx,
                end_column=6,
            )
            c = ws.cell(row=row_idx, column=1)
            c.value = self.processar_cabecalho_html(self.context["cabecalho"])
            c.font = label_font
            c.alignment = left
            row_idx += 1
        row_idx += 1
        headers = [
            "Ordem de Escolha",
            "Inscrição",
            "Nome",
            "Class. Geral",
            "Class. PcD",
            "Class. NNA",
        ]
        for cargo in cargos_list or []:
            cargo_nome = cargo.get("cargo_nome", "")
            sessoes = cargo.get("sessoes", []) or []
            for sessao in sessoes:
                numero_sessao = sessao.get("numero_sessao", "")
                horario_formatado = sessao.get("horario_formatado", "")
                sessao_texto = f"{numero_sessao}ª SESSÃO"
                if horario_formatado:
                    sessao_texto += f" - Horário: {horario_formatado}"
                ws.merge_cells(
                    start_row=row_idx,
                    start_column=1,
                    end_row=row_idx,
                    end_column=6,
                )
                c = ws.cell(row=row_idx, column=1)
                c.value = sessao_texto
                c.font = label_font
                c.alignment = left
                row_idx += 1
                ws.merge_cells(
                    start_row=row_idx,
                    start_column=1,
                    end_row=row_idx,
                    end_column=6,
                )
                c = ws.cell(row=row_idx, column=1)
                c.value = f"CARGO: {cargo_nome}"
                c.font = label_font
                c.alignment = left
                row_idx += 1
                for col, h in enumerate(headers, start=1):
                    cell = ws.cell(row=row_idx, column=col)
                    cell.value = h
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = center if col in (1, 2, 4, 5, 6) else left
                    cell.border = border
                for r, cand in enumerate(
                    sessao.get("candidatos", []), start=row_idx + 1
                ):
                    values = [
                        cand.get("ordem_escolha"),
                        cand.get("codigo_inscricao") or cand.get("uuid"),
                    ]
                    nome = ""
                    status_especial = cand.get("status_especial", "")
                    if status_especial:
                        nome = status_especial
                    else:
                        candidato_obj = cand.get("candidato", {})
                        nome = (
                            candidato_obj.get("nome", "N/A")
                            if isinstance(candidato_obj, dict)
                            else "N/A"
                        )
                        categoria_efetiva = cand.get("categoria_efetiva", "")
                        if categoria_efetiva == "NNA" and cand.get(
                            "classificacao_nna"
                        ):
                            nome += " (NNA)"
                        elif categoria_efetiva == "PCD" and cand.get(
                            "classificacao_pcd"
                        ):
                            nome += " (PCD)"
                    values.append(nome)
                    values.extend(
                        [
                            "-"
                            if str(cand.get("classificacao", "-")) == "None"
                            else cand.get("classificacao", "-"),
                            "-"
                            if str(cand.get("classificacao_pcd", "-"))
                            == "None"
                            else cand.get("classificacao_pcd", "-"),
                            "-"
                            if str(cand.get("classificacao_nna", "-"))
                            == "None"
                            else cand.get("classificacao_nna", "-"),
                        ]
                    )
                    for col, val in enumerate(values, start=1):
                        cell = ws.cell(row=r, column=col)
                        cell.value = val
                        cell.font = normal_font
                        cell.alignment = (
                            center if col in (1, 2, 4, 5, 6) else left
                        )
                        cell.border = border
                row_idx = row_idx + 1 + len(sessao.get("candidatos", [])) + 1
        ws.column_dimensions["A"].width = 18
        ws.column_dimensions["B"].width = 18
        ws.column_dimensions["C"].width = 45
        ws.column_dimensions["D"].width = 16
        ws.column_dimensions["E"].width = 16
        ws.column_dimensions["F"].width = 16
        texto_final = (
            (context or self.context).get("texto_final")
            if context or self.context
            else ""
        )
        if texto_final:
            ws.merge_cells(
                start_row=row_idx,
                start_column=1,
                end_row=row_idx,
                end_column=6,
            )
            c = ws.cell(row=row_idx, column=1)
            c.value = self.processar_cabecalho_html(texto_final)
            c.font = normal_font
            c.alignment = Alignment(
                horizontal="left", vertical="top", wrap_text=True
            )
        from io import BytesIO

        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        for p in temp_image_paths:
            try:
                if os.path.exists(p):
                    os.unlink(p)
            except Exception:
                pass
        resp = HttpResponse(
            buf.read(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )  # type: ignore[assignment]
        resp["Content-Disposition"] = f'attachment; filename="{filename}"'  # type: ignore[index]
        return resp

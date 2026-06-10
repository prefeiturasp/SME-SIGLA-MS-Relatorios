"""Módulo services/relatorios/lista_candidatos_sessao."""

from __future__ import annotations

import logging
import os
import tempfile
from typing import Any

import requests
from django.conf import settings
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render

from relatorios.services.agendas_api_service import AgendasService
from relatorios.services.base.relatorio_base import RelatorioBase
from relatorios.services.candidatos_api_service import CandidatosService

logger = logging.getLogger(__name__)
try:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    OPENPYXL_AVAILABLE = True
except Exception:
    OPENPYXL_AVAILABLE = False
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


class ListaCandidatosSessao(RelatorioBase):
    """Gera relatório de lista de candidatos por sessão, a partir de uma."""

    TEMPLATE_NAME = "relatorios/lista_candidatos_sessao.html"

    def __init__(self, **kwargs: Any) -> None:
        """Inicializa a instância com os parâmetros informados.

        Args:
            self: Instância do objeto.
            **kwargs: Argumentos nomeados variáveis.
        """
        super().__init__(**kwargs)
        self.candidatos_service = CandidatosService(
            base_url=settings.CANDIDATOS_API_URL
        )
        self.agendas_service = AgendasService(
            base_url=settings.AGENDAS_API_URL
        )

    def _fetch_candidatos(
        self, candidatos_uuids: list[str], order_by: str = "ranking_escolha"
    ) -> list[dict[str, Any]]:
        """Fetch candidatos.

        Args:
            self: Instância do objeto.
            candidatos_uuids: Candidatos uuids utilizado na operação.
            order_by: Order by utilizado na operação.

        Returns:
            Lista com os registros obtidos.
        """
        if not candidatos_uuids:
            return []
        resp = self.candidatos_service.buscar_por_uuids(
            uuids=candidatos_uuids, order_by=order_by
        )
        data = resp.json()
        if isinstance(data, dict) and "results" in data:
            return data.get("results", [])  # type: ignore[no-any-return]
        if isinstance(data, list):
            return data
        return []

    @staticmethod
    def _flatten_candidato(item: dict[str, Any]) -> dict[str, Any]:
        """Flatten candidato.

        Args:
            item: Item utilizado na operação.

        Returns:
            Dicionário com os dados retornados pela operação.
        """
        cand = item.get("candidato") or {}
        return {
            "classificacao": item.get("classificacao"),
            "classificacao_nna": item.get("classificacao_nna"),
            "classificacao_pcd": item.get("classificacao_pcd"),
            "inscricao": item.get("codigo_inscricao") or item.get("inscricao"),
            "nome": cand.get("nome") or item.get("nome"),
            "cpf": cand.get("cpf") or item.get("cpf"),
        }

    def _build_context(
        self, candidatos: list[dict[str, Any]], agenda_data: dict[str, Any]
    ) -> dict[str, Any]:
        """Monta context.

        Args:
            self: Instância do objeto.
            candidatos: Candidatos utilizado na operação.
            agenda_data: Agenda data utilizado na operação.

        Returns:
            Dicionário com os dados retornados pela operação.
        """
        linhas = [self._flatten_candidato(c) for c in candidatos]
        return {
            "candidatos": linhas,
            "agenda": agenda_data,
            "agendas": [{"agenda": agenda_data, "candidatos": linhas}],
        }

    def _render_xls(
        self,
        context: dict[str, Any],
        filename: str = "lista_candidatos_sessao.xlsx",
    ) -> HttpResponse:
        """Render xls.

        Args:
            self: Instância do objeto.
            context: Contexto de serialização ou renderização.
            filename: Filename utilizado na operação.

        Returns:
            Resposta HTTP com o resultado da operação.

        Raises:
            ImportError: Se ocorrer erro nesta operação.
        """
        if not OPENPYXL_AVAILABLE:
            raise ImportError(
                "openpyxl não está instalado. Instale com: pip install openpyxl>=3.1.0"  # noqa: E501
            )
        wb = Workbook()
        ws = wb.active
        ws.title = "Candidatos"
        header_fill = PatternFill(
            start_color="ECF0F1", end_color="ECF0F1", fill_type="solid"
        )
        header_font = Font(bold=True, size=11)
        normal_font = Font(size=10)
        center = Alignment(horizontal="center", vertical="center")
        left = Alignment(horizontal="left", vertical="center")
        border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )
        row_idx = 1
        temp_image_paths = []
        logo_url = (
            (context or self.context).get("logo_url")
            if context or self.context
            else ""
        )
        if context.get("usar_logotipo") and logo_url:
            image_path = None
            try:
                if logo_url.startswith("http://") or logo_url.startswith(
                    "https://"
                ):
                    with tempfile.NamedTemporaryFile(
                        suffix=".png", delete=False
                    ) as tmpf:
                        resp = requests.get(logo_url, timeout=15)
                        resp.raise_for_status()
                        tmpf.write(resp.content)
                        image_path = tmpf.name
                        temp_image_paths.append(image_path)
                elif os.path.exists(logo_url):
                    image_path = logo_url
                if image_path:
                    img = XLImage(image_path)
                    try:
                        img.width = 220
                        img.height = 90
                    except Exception:
                        pass
                    ws.add_image(img, "B1")
                    row_idx = max(row_idx, 8)
            except Exception as exc:
                logger.warning(
                    "Não foi possível inserir o logotipo no XLS (lista_candidatos_sessao): %s",  # noqa: E501
                    exc,
                )
        cabecalho_padrao = self.context.get("cabecalho_padrao", "")
        if cabecalho_padrao:
            ws.merge_cells(
                start_row=row_idx,
                start_column=1,
                end_row=row_idx,
                end_column=6,
            )
            title_cell = ws.cell(row=row_idx, column=1)
            title_cell.value = self.processar_cabecalho_html(cabecalho_padrao)
            title_cell.font = Font(bold=True, size=14)
            title_cell.alignment = center
            row_idx += 2
        if self.context.get("cabecalho"):
            ws.merge_cells(
                start_row=row_idx,
                start_column=1,
                end_row=row_idx,
                end_column=6,
            )
            title_cell = ws.cell(row=row_idx, column=1)
            title_cell.value = self.processar_cabecalho_html(
                self.context["cabecalho"]
            )
            title_cell.font = Font(bold=True, size=14)
            title_cell.alignment = center
            row_idx += 2
        ws.merge_cells(
            start_row=row_idx, start_column=1, end_row=row_idx, end_column=6
        )
        title_cell = ws.cell(row=row_idx, column=1)
        title_cell.value = "Lista de Candidatos por Sessão"
        title_cell.font = Font(bold=True, size=14)
        title_cell.alignment = center
        row_idx += 2

        def _fmt_data(date_str: str) -> str:
            """Fmt data.

            Args:
                date_str: Date str utilizado na operação.

            Returns:
                Texto resultante da operação.
            """
            return (
                f"{date_str[8:10]}/{date_str[5:7]}/{date_str[:4]}"
                if len(date_str) >= 10
                else date_str
            )

        def _fmt_hora(time_str: str) -> str:
            """Fmt hora.

            Args:
                time_str: Time str utilizado na operação.

            Returns:
                Texto resultante da operação.
            """
            return time_str[:5] if len(time_str) >= 5 else time_str

        sections = context.get("agendas") or []
        if not sections:
            sections = [
                {
                    "agenda": context.get("agenda") or {},
                    "candidatos": context.get("candidatos") or [],
                }
            ]
        for idx, sec in enumerate(sections):
            agenda = sec.get("agenda") or {}
            escolha_em = agenda.get("escolha_em") or ""
            hora_ini = agenda.get("hora_convocacao_inicio") or ""
            hora_fim = agenda.get("hora_convocacao_fim") or ""
            sessao = agenda.get("sessao") or ""
            cargo_nome = agenda.get("cargo_nome") or ""
            if escolha_em:
                ws.merge_cells(
                    start_row=row_idx,
                    start_column=1,
                    end_row=row_idx,
                    end_column=6,
                )
                c = ws.cell(row=row_idx, column=1)
                c.value = f"Data: {_fmt_data(escolha_em)}"
                c.font = Font(bold=True, size=12)
                c.alignment = left
                row_idx += 1
            if hora_ini or hora_fim:
                ws.merge_cells(
                    start_row=row_idx,
                    start_column=1,
                    end_row=row_idx,
                    end_column=6,
                )
                c = ws.cell(row=row_idx, column=1)
                ini = _fmt_hora(hora_ini) if hora_ini else ""
                fim = _fmt_hora(hora_fim) if hora_fim else ""
                c.value = (
                    f"Horário: {ini} às {fim}"
                    if ini and fim
                    else f"Horário: {ini or fim}"
                )
                c.font = Font(bold=True, size=12)
                c.alignment = left
                row_idx += 1
            if sessao:
                ws.merge_cells(
                    start_row=row_idx,
                    start_column=1,
                    end_row=row_idx,
                    end_column=6,
                )
                c = ws.cell(row=row_idx, column=1)
                c.value = str(sessao)
                c.font = Font(bold=True, size=12)
                c.alignment = left
                row_idx += 1
            if cargo_nome:
                ws.merge_cells(
                    start_row=row_idx,
                    start_column=1,
                    end_row=row_idx,
                    end_column=6,
                )
                c = ws.cell(row=row_idx, column=1)
                c.value = f"Cargo: {cargo_nome}"
                c.font = Font(bold=True, size=12)
                c.alignment = left
                row_idx += 1
            if idx == 0 or escolha_em or hora_ini or hora_fim or sessao:
                row_idx += 1
            headers = [
                "Classificação",
                "Classificação NNA",
                "Classificação PCD",
                "Inscrição",
                "Nome",
                "CPF",
            ]
            for col, h in enumerate(headers, start=1):
                cell = ws.cell(row=row_idx, column=col)
                cell.value = h
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = center
                cell.border = border
            for i, row in enumerate(
                sec.get("candidatos", []), start=row_idx + 1
            ):
                values = [
                    row.get("classificacao"),
                    row.get("classificacao_nna"),
                    row.get("classificacao_pcd"),
                    row.get("inscricao"),
                    row.get("nome"),
                    row.get("cpf"),
                ]
                for col, val in enumerate(values, start=1):
                    cell = ws.cell(row=i, column=col)
                    cell.value = val
                    cell.font = normal_font
                    cell.alignment = center if col in (1, 2, 3) else left
                    cell.border = border
            row_idx = row_idx + 1 + len(sec.get("candidatos", [])) + 1
        ws.column_dimensions["A"].width = 16
        ws.column_dimensions["B"].width = 22
        ws.column_dimensions["C"].width = 22
        ws.column_dimensions["D"].width = 16
        ws.column_dimensions["E"].width = 40
        ws.column_dimensions["F"].width = 20
        if context.get("texto_final"):
            row_idx += 1
            ws.merge_cells(f"A{row_idx}:C{row_idx}")
            cell = ws[f"A{row_idx}"]
            cell.value = self.processar_cabecalho_html(
                context.get("texto_final")
            )  # type: ignore[arg-type]
            cell.font = normal_font
            cell.alignment = Alignment(
                horizontal="left", vertical="top", wrap_text=True
            )
        import io

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        for p in temp_image_paths:
            try:
                if os.path.exists(p):
                    os.unlink(p)
            except Exception:
                pass
        resp = HttpResponse(
            buf.read(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )  # type: ignore[assignment]
        resp["Content-Disposition"] = f'attachment; filename="{filename}"'  # type: ignore[index]
        return resp  # type: ignore[return-value]

    def _render_docx(
        self,
        context: dict[str, Any],
        filename: str = "lista_candidatos_sessao.docx",
    ) -> HttpResponse:
        """Render docx.

        Args:
            self: Instância do objeto.
            context: Contexto de serialização ou renderização.
            filename: Filename utilizado na operação.

        Returns:
            Resposta HTTP com o resultado da operação.

        Raises:
            ImportError: Se ocorrer erro nesta operação.
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx não está instalado. Instale com: pip install python-docx>=0.8.11"  # noqa: E501
            )
        doc = Document()
        doc_sections = doc.sections
        for section in doc_sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        for cab in [
            self.context.get("cabecalho_padrao", ""),
            self.context.get("cabecalho", ""),
        ]:
            if cab:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(self.processar_cabecalho_html(cab))
                run.font.size = Pt(14)
                run.font.bold = True
                doc.add_paragraph()

        def _fmt_data(date_str: str) -> str:
            """Fmt data.

            Args:
                date_str: Date str utilizado na operação.

            Returns:
                Texto resultante da operação.
            """
            return (
                f"{date_str[8:10]}/{date_str[5:7]}/{date_str[:4]}"
                if len(date_str) >= 10
                else date_str
            )

        def _fmt_hora(time_str: str) -> str:
            """Fmt hora.

            Args:
                time_str: Time str utilizado na operação.

            Returns:
                Texto resultante da operação.
            """
            return time_str[:5] if len(time_str) >= 5 else time_str

        sections_list = context.get("agendas") or []
        if not sections_list:
            sections_list = [
                {
                    "agenda": context.get("agenda") or {},
                    "candidatos": context.get("candidatos") or [],
                }
            ]
        for idx, sec in enumerate(sections_list):
            agenda = sec.get("agenda") or {}
            escolha_em = agenda.get("escolha_em") or ""
            hora_ini = agenda.get("hora_convocacao_inicio") or ""
            hora_fim = agenda.get("hora_convocacao_fim") or ""
            sessao = agenda.get("sessao") or ""
            cargo_nome = agenda.get("cargo_nome") or ""
            if escolha_em:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(f"Data: {_fmt_data(escolha_em)}")
                run.font.size = Pt(11)
                run.font.bold = True
            if hora_ini or hora_fim:
                ini = _fmt_hora(hora_ini) if hora_ini else ""
                fim = _fmt_hora(hora_fim) if hora_fim else ""
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(
                    f"Horário: {ini} às {fim}"
                    if ini and fim
                    else f"Horário: {ini or fim}"
                )
                run.font.size = Pt(11)
                run.font.bold = True
            if sessao:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(str(sessao))
                run.font.size = Pt(11)
                run.font.bold = True
            if cargo_nome:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(f"Cargo: {cargo_nome}")
                run.font.size = Pt(11)
                run.font.bold = True
            doc.add_paragraph()
            rows = len(sec.get("candidatos", [])) + 1
            table = doc.add_table(rows=rows, cols=6)
            table.style = "Light Grid Accent 1"
            headers = [
                "Classificação",
                "Classificação NNA",
                "Classificação PCD",
                "Inscrição",
                "Nome",
                "CPF",
            ]
            hdr_cells = table.rows[0].cells
            for j, h in enumerate(headers):
                cell = hdr_cells[j]
                cell.text = h
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                tc_pr = cell._element.get_or_add_tcPr()
                existing_shd = tc_pr.find(qn("w:shd"))
                if existing_shd is not None:
                    tc_pr.remove(existing_shd)
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:fill"), "ECF0F1")
                shading_elm.set(qn("w:val"), "clear")
                tc_pr.append(shading_elm)
            for i, row in enumerate(sec.get("candidatos", []), start=1):
                cells = table.rows[i].cells
                values = [
                    str(row.get("classificacao") or ""),
                    str(row.get("classificacao_nna") or ""),
                    str(row.get("classificacao_pcd") or ""),
                    str(row.get("inscricao") or ""),
                    str(row.get("nome") or ""),
                    str(row.get("cpf") or ""),
                ]
                for col_idx, val in enumerate(values):
                    cell = cells[col_idx]
                    cell.text = val
                    if col_idx in (0, 1, 2):
                        cell.paragraphs[
                            0
                        ].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
            if idx < len(sections_list) - 1:
                doc.add_paragraph()
        if context.get("texto_final"):
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(
                self.processar_cabecalho_html(context.get("texto_final"))
            )  # type: ignore[arg-type]
            run.font.size = Pt(10)
            doc.add_paragraph()
        import io

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        resp = HttpResponse(
            buf.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        resp["Content-Disposition"] = f'attachment; filename="{filename}"'
        return resp

    def gerar(
        self,
        processo_uuid: str,
        request: Any,
        formato: str = "html",
        cabecalho: str = "",
        agenda_uuid: str = "",
        **kwargs: Any,
    ) -> tuple[HttpResponse, dict[str, Any]]:  # type: ignore[override]
        """Gera a lista de candidatos por sessão a partir de UUIDs.

        Args:
            self: Instância do objeto.
            processo_uuid: UUID do processo de convocação.
            request: Requisição HTTP recebida.
            formato: Formato utilizado na operação.
            cabecalho: Cabecalho utilizado na operação.
            agenda_uuid: UUID de agenda.
            **kwargs: Argumentos nomeados variáveis.

        Returns:
            Tupla com os objetos criados ou atualizados.
        """
        try:
            if agenda_uuid:
                agenda_resp = self.agendas_service.buscar_agenda_por_uuid(
                    str(agenda_uuid)
                )
            else:
                agenda_resp = self.agendas_service.buscar_agendas(
                    processo_convocacao_uuid=str(processo_uuid)
                )
            raw = agenda_resp.json()
            if (
                isinstance(raw, dict)
                and "results" in raw
                and isinstance(raw["results"], list)
            ):
                agendas_list: list[dict[str, Any]] = raw["results"]
            elif isinstance(raw, list):
                agendas_list = [a for a in raw if isinstance(a, dict)]
            elif isinstance(raw, dict):
                agendas_list = [raw]
            else:
                agendas_list = []
            sections: list[dict[str, Any]] = []
            for a in agendas_list:
                if a.get("retardatario") is not False:
                    continue
                uuids = a.get("candidatos_uuids") or []
                uuids = [u for u in uuids if isinstance(u, str)]
                cand_list = self._fetch_candidatos(uuids)
                linhas = [self._flatten_candidato(c) for c in cand_list]
                sections.append({"agenda": a, "candidatos": linhas})
            if not sections:
                context = {"agendas": [], "agenda": {}, "candidatos": []}  # type: ignore[var-annotated]
            elif len(sections) == 1:
                context = {
                    "agendas": sections,
                    "agenda": sections[0]["agenda"],
                    "candidatos": sections[0]["candidatos"],
                }
            else:
                context = {"agendas": sections}
            if cabecalho is not None:
                self.context["cabecalho"] = cabecalho
            logo_url = (
                request.build_absolute_uri(self.context.get("logo_url", ""))
                if self.context.get("logo_url")
                else ""
            )
            self.context["is_pdf"] = False
            self.context["logo_url"] = logo_url
            self.context.update(context)
        except Exception as exc:
            logger.error(
                "Erro ao processar agenda/candidatos: %s", exc, exc_info=True
            )
            raise
        if formato == "pdf":
            self.context.update({"is_pdf": True})
            logger.info("Gerando PDF lista_candidatos_sessao")
            return (
                self.render_to_pdf(
                    self.TEMPLATE_NAME,
                    self.context,
                    filename="lista_candidatos_sessao.pdf",
                ),
                context,
            )
        if formato == "html":
            logger.info("Gerando HTML lista_candidatos_sessao")
            return (render(request, self.TEMPLATE_NAME, self.context), context)
        if formato in ("xls", "xlsx"):
            logger.info("Gerando XLS lista_candidatos_sessao")
            return (self._render_xls(self.context), self.context)
        if formato in ("doc", "docx"):
            logger.info("Gerando DOCX lista_candidatos_sessao")
            return (self._render_docx(self.context), self.context)
        return (JsonResponse(self.context, safe=False), self.context)

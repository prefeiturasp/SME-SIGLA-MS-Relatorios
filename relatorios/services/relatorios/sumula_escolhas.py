"""Módulo services/relatorios/sumula_escolhas."""

from __future__ import annotations

import logging
import os
import tempfile
from io import BytesIO
from typing import Any

import requests
from django.conf import settings
from django.http import HttpResponse
from django.shortcuts import render

from relatorios.services.base.relatorio_base import RelatorioBase
from relatorios.services.candidatos_api_service import CandidatosService
from relatorios.services.escolhas_api_service import EscolhasService
from relatorios.services.processos_api_service import ProcessosService

try:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
logger = logging.getLogger(__name__)


class SumulaEscolhas(RelatorioBase):
    """Classe concreta responsável por gerar o relatório de Súmula de."""

    TEMPLATE_NAME = "relatorios/sumula_escolhas.html"

    def __init__(self, **kwargs: Any) -> None:
        """Inicializa a instância com os parâmetros informados.

        Args:
            self: Instância do objeto.
            **kwargs: Argumentos nomeados variáveis.
        """
        super().__init__(**kwargs)
        self.escolhas_service = EscolhasService(
            base_url=settings.ESCOLHAS_API_URL
        )
        self.candidatos_service = CandidatosService(
            base_url=settings.CANDIDATOS_API_URL
        )
        self.processos_service = ProcessosService(
            base_url=settings.PROCESSOS_API_URL
        )

    def gerar(
        self,
        processo_uuid: str,
        request: Any,
        formato: str = "html",
        cabecalho: str = "",
        **kwargs: Any,
    ) -> Any:
        """Gera o relatório de Súmula de Escolhas.

        Args:
            self: Instância do objeto.
            processo_uuid: UUID do processo de convocação.
            request: Requisição HTTP recebida.
            formato: Formato utilizado na operação.
            cabecalho: Cabecalho utilizado na operação.
            **kwargs: Argumentos nomeados variáveis.

        Returns:
            Valor calculado conforme a regra aplicada.
        """
        cargos_map = {}
        try:
            cargos_response = (
                self.processos_service.buscar_cargos_por_processo(
                    processo_uuid=str(processo_uuid) if processo_uuid else ""
                )
            )
            cargos_data = cargos_response.json()
            cargos = cargos_data if isinstance(cargos_data, list) else []
            for cargo in cargos:
                codigo = (
                    cargo.get("cargo_codigo")
                    or cargo.get("codigo_cargo")
                    or ""
                )
                nome = cargo.get("cargo_nome") or cargo.get("nome") or ""
                if codigo and nome:
                    cargos_map[str(codigo)] = nome
                    if isinstance(codigo, int | float):
                        cargos_map[codigo] = nome  # type: ignore[index]
        except Exception as exc:
            logger.warning(
                "Falha ao buscar cargos do processo: %s. Continuando sem mapeamento de cargos.",  # noqa: E501
                exc,
            )
        try:
            candidatos_response = self.candidatos_service.buscar_concurso_candidatos_por_processo(  # noqa: E501
                processo_uuid=str(processo_uuid) if processo_uuid else ""
            )
            candidatos_data = candidatos_response.json()
            candidatos = (
                candidatos_data.get("results", [])
                if isinstance(candidatos_data, dict)
                else candidatos_data
            )
        except Exception as exc:
            logger.error("Falha ao buscar candidatos da API externa: %s", exc)
            raise
        candidatos_map = {}
        for candidato in candidatos:
            candidato_uuid = candidato.get("uuid")
            if candidato_uuid:
                candidatos_map[str(candidato_uuid)] = candidato
        concurso_candidato_uuids = [
            candidato.get("uuid")
            for candidato in candidatos
            if candidato.get("uuid")
        ]
        try:
            escolhas_data = (
                self.escolhas_service.buscar_escolhas_por_candidatos(
                    candidato_uuids=concurso_candidato_uuids, situacao=None
                )
            )  # type: ignore[arg-type]
        except Exception as exc:
            logger.error("Falha ao buscar escolhas da API externa: %s", exc)
            raise
        escolhas_realizadas = [
            e
            for e in escolhas_data
            if e.get("situacao") not in ["nao-escolha", "reconvocacao", None]
        ]
        escolhas_com_candidatos = []
        for escolha in escolhas_realizadas:
            candidato_uuid = escolha.get("candidato_uuid")
            if not candidato_uuid:
                continue
            candidato = candidatos_map.get(
                str(candidato_uuid)
            ) or candidatos_map.get(candidato_uuid)
            if not candidato:
                continue
            candidato_obj = (
                candidato.get("candidato", {})
                if isinstance(candidato.get("candidato"), dict)
                else {}
            )
            classificacao_geral = candidato.get("classificacao")
            classificacao_nna = candidato.get("classificacao_nna")
            classificacao_pcd = candidato.get("classificacao_pcd")
            categoria_efetiva = (
                candidato.get("categoria_efetiva") or ""
            ).upper()
            classificacao_coluna_geral = (
                classificacao_geral if classificacao_geral is not None else "-"
            )
            classificacao_coluna_nna = "-"
            classificacao_coluna_pcd = "-"
            if categoria_efetiva == "NNA":
                classificacao_coluna_nna = (
                    classificacao_nna if classificacao_nna is not None else "-"
                )
                classificacao_ordem = classificacao_nna
            elif categoria_efetiva == "PCD":
                classificacao_coluna_pcd = (
                    classificacao_pcd if classificacao_pcd is not None else "-"
                )
                classificacao_ordem = classificacao_pcd
            else:
                classificacao_ordem = classificacao_geral
            nome = candidato_obj.get("nome") or "-"
            tipo_candidato = None
            if tipo_candidato:
                nome_formatado = f"{nome} - {tipo_candidato}"
            else:
                nome_formatado = nome
            vaga_escola = escolha.get("vaga_escola", {})
            escola = (
                vaga_escola.get("escola", {})
                if isinstance(vaga_escola, dict)
                else {}
            )
            dre = escola.get("dre", {}) if isinstance(escola, dict) else {}
            cargo_codigo = candidato.get("codigo_cargo") or ""
            cargo_descricao = candidato.get("descricao_cargo") or ""
            if not cargo_descricao and cargo_codigo:
                cargo_descricao = (
                    cargos_map.get(str(cargo_codigo))
                    or cargos_map.get(cargo_codigo)
                    or ""
                )
            if not cargo_descricao and cargo_codigo:
                cargo_descricao = f"Cargo {cargo_codigo}"
            elif not cargo_descricao:
                cargo_descricao = "Cargo não informado"
            nome_oficial = (
                escola.get("nome_oficial", "-")
                if isinstance(escola, dict)
                else "-"
            )
            codigo_eol = (
                escola.get("codigo_eol", "")
                if isinstance(escola, dict)
                else ""
            )
            dre_nome = dre.get("nome", "-") if isinstance(dre, dict) else "-"
            dre_codigo = dre.get("codigo", "") if isinstance(dre, dict) else ""
            tipo_vaga_raw = escolha.get("tipo_vaga", "")
            if tipo_vaga_raw == "precaria":
                tipo_vaga = "P"
            elif tipo_vaga_raw == "definitiva":
                tipo_vaga = "D"
            else:
                tipo_vaga = "-"
            escolhas_com_candidatos.append(
                {
                    "cargo_codigo": cargo_codigo,
                    "cargo_descricao": cargo_descricao,
                    "dre_codigo": dre_codigo,
                    "dre_nome": dre_nome,
                    "escola_nome": nome_oficial,
                    "escola_codigo_eol": codigo_eol,
                    "classificacao": classificacao_coluna_geral,
                    "classificacao_nna": classificacao_coluna_nna,
                    "classificacao_pcd": classificacao_coluna_pcd,
                    "classificacao_ordem": classificacao_ordem,
                    "nome_candidato": nome_formatado,
                    "tipo_vaga": tipo_vaga,
                }
            )
        cargos_list = self._agrupar_por_cargo_dre_e_escola(
            escolhas_com_candidatos
        )
        logo_url = (
            request.build_absolute_uri(self.context.get("logo_url", ""))
            if self.context.get("logo_url")
            else ""
        )
        self.context.update(
            {"cargos": cargos_list, "is_pdf": False, "logo_url": logo_url}
        )
        if formato == "xls" or formato == "csv":
            filename = f"relatorio_sumula_escolhas_{processo_uuid}.xlsx"
            logger.info("Gerando Excel: %s", filename)
            response = self.render_to_xls(
                context=self.context, filename=filename
            )
            return (response, cargos_list)
        elif formato == "docx" or formato == "doc":
            filename = f"relatorio_sumula_escolhas_{processo_uuid}.docx"
            logger.info("Gerando Word: %s", filename)
            cabecalho_docx = (
                cabecalho.strip()
                if cabecalho and cabecalho.strip()
                else self.context.get("cabecalho_padrao", "")
            )
            response = self.render_to_docx(
                cargos_list, cabecalho_docx, filename=filename
            )
            return (response, cargos_list)
        elif formato == "pdf":
            filename = f"relatorio_sumula_escolhas_{processo_uuid}.pdf"
            logger.info("Gerando PDF: %s", filename)
            self.context.update({"is_pdf": True, "cargos": cargos_list})
            response = self.render_to_pdf(
                self.TEMPLATE_NAME, self.context, filename=filename
            )
            return (response, cargos_list)
        else:
            logger.info("Gerando HTML")
            self.context.update({"cargos": cargos_list})
            response = render(request, self.TEMPLATE_NAME, self.context)
            return (response, cargos_list)

    def _agrupar_por_cargo_dre_e_escola(self, escolhas: list) -> list:
        """Agrupa escolhas por cargo, depois por DRE e depois por Unidade.

        Args:
            self: Instância do objeto.
            escolhas: Escolhas utilizado na operação.

        Returns:
            Lista com os registros obtidos.
        """
        cargos_dict = {}  # type: ignore[var-annotated]
        for escolha in escolhas:
            cargo_codigo = escolha.get("cargo_codigo", "") or ""
            cargo_descricao = escolha.get("cargo_descricao", "") or ""
            dre_codigo = escolha.get("dre_codigo", "") or ""
            dre_nome = escolha.get("dre_nome", "") or "-"
            escola_nome = escolha.get("escola_nome", "") or "-"
            escola_codigo_eol = escolha.get("escola_codigo_eol", "") or ""
            if not cargo_descricao:
                if cargo_codigo and cargo_codigo != "-":
                    cargo_descricao = f"Cargo {cargo_codigo}"
                else:
                    cargo_descricao = "Cargo não informado"
            if not dre_nome or dre_nome == "-":
                if dre_codigo:
                    dre_nome = f"DRE {dre_codigo}"
                else:
                    dre_nome = "DRE não informada"
            if not escola_nome or escola_nome == "-":
                if escola_codigo_eol:
                    escola_nome = f"Escola EOL {escola_codigo_eol}"
                else:
                    escola_nome = "Unidade Escolar não informada"
            escola_chave = (
                f"{escola_nome}_{escola_codigo_eol}"
                if escola_codigo_eol
                else escola_nome
            )
            if cargo_codigo not in cargos_dict:
                cargos_dict[cargo_codigo] = {
                    "codigo": cargo_codigo
                    if cargo_codigo and cargo_codigo != "-"
                    else "",
                    "descricao": cargo_descricao,
                    "dres": {},
                }
            if dre_codigo not in cargos_dict[cargo_codigo]["dres"]:
                cargos_dict[cargo_codigo]["dres"][dre_codigo] = {
                    "codigo": dre_codigo if dre_codigo else "",
                    "nome": dre_nome,
                    "escolas": {},
                }
            if (
                escola_chave
                not in cargos_dict[cargo_codigo]["dres"][dre_codigo]["escolas"]
            ):
                cargos_dict[cargo_codigo]["dres"][dre_codigo]["escolas"][
                    escola_chave
                ] = {
                    "nome": escola_nome,
                    "codigo_eol": escola_codigo_eol,
                    "escolhas": [],
                }
            cargos_dict[cargo_codigo]["dres"][dre_codigo]["escolas"][
                escola_chave
            ]["escolhas"].append(escolha)
        cargos_list = []
        for cargo_codigo, cargo_data in cargos_dict.items():
            dres_list = []
            for dre_codigo, dre_data in cargo_data["dres"].items():
                escolas_list = []
                for escola_chave, escola_data in dre_data["escolas"].items():
                    escola_data["escolhas"].sort(
                        key=lambda e: e.get("classificacao_ordem")
                        if isinstance(
                            e.get("classificacao_ordem"), int | float
                        )
                        else e.get("classificacao")
                        if isinstance(e.get("classificacao"), int | float)
                        else float("inf")
                    )
                    escolas_list.append(escola_data)
                escolas_list.sort(key=lambda e: e["nome"])
                dres_list.append(
                    {
                        "codigo": dre_data["codigo"],
                        "nome": dre_data["nome"],
                        "escolas": escolas_list,
                    }
                )
            dres_list.sort(key=lambda d: d["nome"])
            cargos_list.append(
                {
                    "codigo": cargo_data["codigo"],
                    "descricao": cargo_data["descricao"],
                    "dres": dres_list,
                }
            )
        cargos_list.sort(key=lambda x: x["descricao"])
        return cargos_list

    def render_to_xls(
        self,
        context: Any = None,
        filename: Any = "relatorio_sumula_escolhas.xlsx",
    ) -> Any:
        """Gera um arquivo Excel (XLSX) mantendo a estrutura hierárquica do.

        Args:
            self: Instância do objeto.
            context: Contexto de serialização ou renderização.
            filename: Filename utilizado na operação.

        Returns:
            Valor calculado conforme a regra aplicada.

        Raises:
            ImportError: Se ocorrer erro nesta operação.
        """
        if context is None:
            context = {}
        if not OPENPYXL_AVAILABLE:
            raise ImportError(
                "openpyxl não está instalado. Instale com: pip install openpyxl>=3.1.0"  # noqa: E501
            )
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Súmula de Escolhas"
            cargo_fill = PatternFill(
                start_color="667eea", end_color="667eea", fill_type="solid"
            )
            dre_fill = PatternFill(
                start_color="34495e", end_color="34495e", fill_type="solid"
            )
            escola_fill = PatternFill(
                start_color="5a6c7d", end_color="5a6c7d", fill_type="solid"
            )
            table_header_fill = PatternFill(
                start_color="ECF0F1", end_color="ECF0F1", fill_type="solid"
            )
            cargo_font = Font(bold=True, color="FFFFFF", size=12)
            dre_font = Font(bold=True, color="FFFFFF", size=11)
            escola_font = Font(bold=True, color="FFFFFF", size=10)
            header_font = Font(bold=True, size=10)
            normal_font = Font(size=10)
            title_font = Font(bold=True, size=14)
            border = Border(
                left=Side(style="thin"),
                right=Side(style="thin"),
                top=Side(style="thin"),
                bottom=Side(style="thin"),
            )
            center_align = Alignment(horizontal="center", vertical="center")
            center_wrap_align = Alignment(
                horizontal="center", vertical="center", wrap_text=True
            )
            left_align = Alignment(horizontal="left", vertical="center")
            row = 1
            temp_image_paths = []
            logo_url = (
                (context or self.context).get("logo_url")
                if context or self.context
                else ""
            )
            if context.get("usar_logotipo") and logo_url:
                image_path = None
                try:
                    if logo_url.startswith("http://") or logo_url.startswith(
                        "https://"
                    ):
                        with tempfile.NamedTemporaryFile(
                            suffix=".png", delete=False
                        ) as tmpf:
                            resp = requests.get(logo_url, timeout=15)
                            resp.raise_for_status()
                            tmpf.write(resp.content)
                            image_path = tmpf.name
                            temp_image_paths.append(image_path)
                    elif os.path.exists(logo_url):
                        image_path = logo_url
                    if image_path:
                        img = XLImage(image_path)
                        try:
                            img.width = 220
                            img.height = 90
                        except Exception:
                            pass
                        ws.add_image(img, "B1")
                        row = max(row, 8)
                except Exception as exc:
                    logger.warning(
                        "Não foi possível inserir o logotipo no XLS: %s", exc
                    )
            cabecalho_padrao = self.context.get("cabecalho_padrao", "")
            if cabecalho_padrao:
                ws.merge_cells(f"A{row}:E{row}")
                cell = ws[f"A{row}"]
                cell.value = self.processar_cabecalho_html(cabecalho_padrao)
                cell.font = title_font
                cell.alignment = center_wrap_align
                row += 2
            if self.context.get("cabecalho"):
                ws.merge_cells(f"A{row}:E{row}")
                cell = ws[f"A{row}"]
                cell.value = self.processar_cabecalho_html(
                    self.context["cabecalho"]
                )
                cell.font = title_font
                cell.alignment = center_wrap_align
                row += 2
            for cargo in context.get("cargos", []):
                cargo_descricao = cargo.get("descricao", "")
                ws.merge_cells(f"A{row}:C{row}")
                cell = ws[f"A{row}"]
                cell.value = f"Cargo: {cargo_descricao}"
                cell.font = cargo_font
                cell.fill = cargo_fill
                cell.alignment = left_align
                row += 1
                for dre in cargo.get("dres", []):
                    dre_nome = dre.get("nome", "")
                    ws.merge_cells(f"A{row}:E{row}")
                    cell = ws[f"A{row}"]
                    cell.value = f"DRE - {dre_nome}"
                    cell.font = dre_font
                    cell.fill = dre_fill
                    cell.alignment = left_align
                    row += 1
                    for escola in dre.get("escolas", []):
                        escola_nome = escola.get("nome", "")
                        ws.merge_cells(f"A{row}:E{row}")
                        cell = ws[f"A{row}"]
                        cell.value = escola_nome
                        cell.font = escola_font
                        cell.fill = escola_fill
                        cell.alignment = left_align
                        row += 1
                        headers = [
                            "Classificação",
                            "Classificação NNA",
                            "Classificação PcD",
                            "Candidatos",
                            "Tipo da Vaga",
                        ]
                        for col, header in enumerate(headers, start=1):
                            cell = ws.cell(row=row, column=col)
                            cell.value = header
                            cell.fill = table_header_fill
                            cell.font = header_font
                            cell.alignment = center_align
                            cell.border = border
                        row += 1
                        for escolha in escola.get("escolhas", []):
                            ws.cell(row=row, column=1).value = escolha.get(
                                "classificacao", "-"
                            )
                            ws.cell(row=row, column=2).value = escolha.get(
                                "classificacao_nna", "-"
                            )
                            ws.cell(row=row, column=3).value = escolha.get(
                                "classificacao_pcd", "-"
                            )
                            ws.cell(row=row, column=4).value = escolha.get(
                                "nome_candidato", "-"
                            )
                            ws.cell(row=row, column=5).value = escolha.get(
                                "tipo_vaga", "-"
                            )
                            for col in range(1, 6):
                                cell = ws.cell(row=row, column=col)
                                cell.border = border
                                cell.font = normal_font
                                if col in [1, 2, 3, 5]:
                                    cell.alignment = center_align
                                else:
                                    cell.alignment = left_align
                            row += 1
                        row += 1
                    row += 1
                row += 1
            column_widths = {"A": 15, "B": 18, "C": 18, "D": 50, "E": 15}
            for col_letter, width in column_widths.items():
                ws.column_dimensions[col_letter].width = width
            texto_final = self.context.get("texto_final")
            if texto_final:
                row += 1
                ws.merge_cells(f"A{row}:E{row}")
                cell = ws[f"A{row}"]
                cell.value = self.processar_cabecalho_html(texto_final)
                cell.font = normal_font
                cell.alignment = Alignment(
                    horizontal="left", vertical="top", wrap_text=True
                )
            buffer = BytesIO()
            wb.save(buffer)
            buffer.seek(0)
            for p in temp_image_paths:
                try:
                    if os.path.exists(p):
                        os.unlink(p)
                except Exception:
                    pass
            response = HttpResponse(
                buffer.read(),
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            response["Content-Disposition"] = (
                f'attachment; filename="{filename}"'
            )
            return response
        except Exception as exc:
            logger.error("Erro ao gerar Excel: %s", exc, exc_info=True)
            raise

    def render_to_docx(
        self,
        cargos_list: Any,
        cabecalho: Any,
        filename: Any = "relatorio_sumula_escolhas.docx",
    ) -> Any:
        """Gera um arquivo Word (DOCX) mantendo a estrutura hierárquica do.

        Args:
            self: Instância do objeto.
            cargos_list: Lista de cargos do processo.
            cabecalho: Cabecalho utilizado na operação.
            filename: Filename utilizado na operação.

        Returns:
            Valor calculado conforme a regra aplicada.

        Raises:
            ImportError: Se ocorrer erro nesta operação.
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx não está instalado. Instale com: pip install python-docx>=1.1.0"  # noqa: E501
            )
        texto_final = self.context.get("texto_final", "")
        try:
            doc = Document()
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            RGBColor(102, 126, 234)
            RGBColor(52, 73, 94)
            RGBColor(90, 108, 125)
            RGBColor(236, 240, 241)
            if cabecalho:
                cabecalho_texto = self.processar_cabecalho_html(cabecalho)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cabecalho_texto)
                run.font.size = Pt(14)
                run.font.bold = True
                doc.add_paragraph()
            for cargo in cargos_list:
                cargo_descricao = cargo.get("descricao", "")
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(f"Cargo: {cargo_descricao}")
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                p_pr = p._element.get_or_add_pPr()
                existing_shd = p_pr.find(qn("w:shd"))
                if existing_shd is not None:
                    p_pr.remove(existing_shd)
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:fill"), "667eea")
                shading_elm.set(qn("w:val"), "clear")
                p_pr.append(shading_elm)
                for dre in cargo.get("dres", []):
                    dre_nome = dre.get("nome", "")
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(f"DRE - {dre_nome}")
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    p_pr = p._element.get_or_add_pPr()
                    existing_shd = p_pr.find(qn("w:shd"))
                    if existing_shd is not None:
                        p_pr.remove(existing_shd)
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:fill"), "34495e")
                    shading_elm.set(qn("w:val"), "clear")
                    p_pr.append(shading_elm)
                    for escola in dre.get("escolas", []):
                        escola_nome = escola.get("nome", "")
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = p.add_run(escola_nome)
                        run.font.size = Pt(10)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        p_pr = p._element.get_or_add_pPr()
                        existing_shd = p_pr.find(qn("w:shd"))
                        if existing_shd is not None:
                            p_pr.remove(existing_shd)
                        shading_elm = OxmlElement("w:shd")
                        shading_elm.set(qn("w:fill"), "5a6c7d")
                        shading_elm.set(qn("w:val"), "clear")
                        p_pr.append(shading_elm)
                        headers = [
                            "Classificação",
                            "Classificação NNA",
                            "Classificação PcD",
                            "Candidatos",
                            "Tipo da Vaga",
                        ]
                        table = doc.add_table(rows=1, cols=len(headers))
                        table.style = "Light Grid Accent 1"
                        header_cells = table.rows[0].cells
                        for i, header in enumerate(headers):
                            cell = header_cells[i]
                            cell.text = header
                            cell.paragraphs[
                                0
                            ].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cell.paragraphs[0].runs[0].font.bold = True
                            cell.paragraphs[0].runs[0].font.size = Pt(10)
                            tc_pr = cell._element.get_or_add_tcPr()
                            existing_shd = tc_pr.find(qn("w:shd"))
                            if existing_shd is not None:
                                tc_pr.remove(existing_shd)
                            shading_elm = OxmlElement("w:shd")
                            shading_elm.set(qn("w:fill"), "ECF0F1")
                            shading_elm.set(qn("w:val"), "clear")
                            tc_pr.append(shading_elm)
                        for escolha in escola.get("escolhas", []):
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(
                                escolha.get("classificacao", "-")
                            )
                            row_cells[1].text = str(
                                escolha.get("classificacao_nna", "-")
                            )
                            row_cells[2].text = str(
                                escolha.get("classificacao_pcd", "-")
                            )
                            row_cells[3].text = str(
                                escolha.get("nome_candidato", "-")
                            )
                            row_cells[4].text = str(
                                escolha.get("tipo_vaga", "-")
                            )
                            row_cells[0].paragraphs[
                                0
                            ].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            row_cells[1].paragraphs[
                                0
                            ].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            row_cells[2].paragraphs[
                                0
                            ].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            row_cells[3].paragraphs[
                                0
                            ].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            row_cells[4].paragraphs[
                                0
                            ].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for cell in row_cells:
                                cell.paragraphs[0].runs[0].font.size = Pt(10)
                        doc.add_paragraph()
                    doc.add_paragraph()
                doc.add_paragraph()
            if texto_final:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(self.processar_cabecalho_html(texto_final))
                run.font.size = Pt(10)
                doc.add_paragraph()
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            response = HttpResponse(
                buffer.read(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = (
                f'attachment; filename="{filename}"'
            )
            return response
        except Exception as exc:
            logger.error("Erro ao gerar Word: %s", exc, exc_info=True)
            raise

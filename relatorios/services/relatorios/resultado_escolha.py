"""Módulo services/relatorios/resultado_escolha."""

from __future__ import annotations

import logging
import os
import re
import tempfile
from io import BytesIO
from typing import Any

import requests
from django.conf import settings
from django.http import HttpResponse
from django.shortcuts import render
from django.utils import timezone

from relatorios.services.agendas_api_service import AgendasService
from relatorios.services.base.relatorio_base import RelatorioBase
from relatorios.services.candidatos_api_service import CandidatosService
from relatorios.services.escolhas_api_service import EscolhasService
from relatorios.services.processos_api_service import ProcessosService

try:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
logger = logging.getLogger(__name__)


class ResultadoEscolha(RelatorioBase):
    """Classe concreta responsável por gerar o relatório de Resultado de."""

    TEMPLATE_NAME = "relatorios/resultado_escolha.html"

    def __init__(self, tipo: str, **kwargs: Any) -> None:
        """Inicializa o service com as dependências necessárias.

        Args:
            self: Instância do objeto.
            tipo: Parâmetro tipo.
            **kwargs: Argumentos nomeados variáveis.

        Raises:
            Nenhuma exceção específica documentada.
        """
        super().__init__(**kwargs)
        self.escolhas_service = EscolhasService(
            base_url=settings.ESCOLHAS_API_URL
        )
        self.candidatos_service = CandidatosService(
            base_url=settings.CANDIDATOS_API_URL
        )
        self.processos_service = ProcessosService(
            base_url=settings.PROCESSOS_API_URL
        )
        self.agendas_service = AgendasService(
            base_url=settings.AGENDAS_API_URL
        )
        self.tipo = tipo

    def gerar(
        self,
        processo_uuid: str,
        request: Any,
        formato: str = "html",
        cabecalho: str = "",
        agenda_uuid: str = None,
        **kwargs: Any,
    ) -> Any:  # type: ignore[assignment]
        """Gera o relatório de Resultado da Escolha SIM.

        Args:
            self: Instância do objeto.
            processo_uuid: UUID do processo de convocação.
            request: Objeto request do Django.
            formato: Formato do relatório ('html', 'pdf', 'xls' ou 'docx').
            cabecalho: Texto do cabeçalho do relatório (opcional).
            agenda_uuid: UUID da agenda.
            **kwargs: Argumentos nomeados variáveis.

        Returns:
            Resultado da operação.

        Raises:
            Nenhuma exceção específica documentada.
        """
        cargos_map = {}
        try:
            cargos_response = (
                self.processos_service.buscar_cargos_por_processo(
                    processo_uuid=str(processo_uuid) if processo_uuid else ""
                )
            )
            cargos_data = cargos_response.json()
            cargos = cargos_data if isinstance(cargos_data, list) else []
            for cargo in cargos:
                codigo = (
                    cargo.get("cargo_codigo")
                    or cargo.get("codigo_cargo")
                    or ""
                )
                nome = cargo.get("cargo_nome") or cargo.get("nome") or ""
                if codigo and nome:
                    cargos_map[str(codigo)] = nome
                    if isinstance(codigo, int | float):
                        cargos_map[codigo] = nome  # type: ignore[index]
        except Exception as exc:
            logger.warning(
                "Falha ao buscar cargos do processo: %s. Continuando sem mapeamento de cargos.",  # noqa: E501
                exc,
            )
        agendas_map = {}
        agendas_por_candidato = {}
        try:
            agendas_response = self.agendas_service.buscar_agendas(
                processo_convocacao_uuid=str(processo_uuid)
                if processo_uuid
                else "",
                page=1,
                page_size=1000,
            )
            agendas_data = agendas_response.json()
            agendas = (
                agendas_data.get("results", [])
                if isinstance(agendas_data, dict)
                else agendas_data
            )
            for agenda in agendas:
                agenda_uuid = agenda.get("uuid")
                if agenda_uuid:
                    agendas_map[str(agenda_uuid)] = agenda
                    candidatos_uuids = agenda.get("candidatos_uuids", [])
                    for candidato_uuid in candidatos_uuids:
                        if candidato_uuid:
                            candidato_uuid_str = str(candidato_uuid)
                            if candidato_uuid_str not in agendas_por_candidato:
                                agendas_por_candidato[candidato_uuid_str] = (
                                    agenda
                                )
        except Exception as exc:
            logger.warning(
                "Falha ao buscar agendas do processo: %s. Continuando sem agendas.",  # noqa: E501
                exc,
            )
        try:
            candidatos_response = self.candidatos_service.buscar_concurso_candidatos_por_processo(  # noqa: E501
                processo_uuid=str(processo_uuid) if processo_uuid else ""
            )
            candidatos_data = candidatos_response.json()
            candidatos = (
                candidatos_data.get("results", [])
                if isinstance(candidatos_data, dict)
                else candidatos_data
            )
        except Exception as exc:
            logger.error("Falha ao buscar candidatos da API externa: %s", exc)
            raise
        candidatos_map = {}
        for candidato in candidatos:
            candidato_uuid = candidato.get("uuid")
            if candidato_uuid:
                candidatos_map[str(candidato_uuid)] = candidato
        concurso_candidato_uuids = [
            candidato.get("uuid")
            for candidato in candidatos
            if candidato.get("uuid")
        ]
        if self.tipo == "RESULTADO_ESCOLHA":
            tipos_escolha = ["escolha", "nao-escolha", "reconvocacao"]
            todas_escolhas = []
            for tipo_escolha in tipos_escolha:
                try:
                    escolhas_data = (
                        self.escolhas_service.buscar_escolhas_por_candidatos(
                            candidato_uuids=concurso_candidato_uuids,
                            situacao=tipo_escolha,
                        )
                    )
                    todas_escolhas.extend(escolhas_data)
                except Exception as exc:
                    logger.warning(
                        "Falha ao buscar escolhas do tipo %s: %s. Continuando...",  # noqa: E501
                        tipo_escolha,
                        exc,
                    )
            escolhas_data = todas_escolhas
        else:
            tipo_escolha = (
                "escolha"
                if self.tipo == "RESULTADO_ESCOLHA_SIM"
                else "nao-escolha"
                if self.tipo == "RESULTADO_ESCOLHA_NAO"
                else "reconvocacao"
                if self.tipo == "RESULTADO_ESCOLHA_RECONVOCACAO"
                else None
            )  # type: ignore[assignment]
            try:
                escolhas_data = (
                    self.escolhas_service.buscar_escolhas_por_candidatos(
                        candidato_uuids=concurso_candidato_uuids,
                        situacao=tipo_escolha,
                    )
                )
            except Exception as exc:
                logger.error(
                    "Falha ao buscar escolhas da API externa: %s", exc
                )
                raise
        escolhas_com_candidatos = []
        for escolha in escolhas_data:
            candidato_uuid = escolha.get("candidato_uuid")
            if not candidato_uuid:
                continue
            candidato = candidatos_map.get(
                str(candidato_uuid)
            ) or candidatos_map.get(candidato_uuid)
            if not candidato:
                continue
            candidato_obj = (
                candidato.get("candidato", {})
                if isinstance(candidato.get("candidato"), dict)
                else {}
            )
            classificacao_geral = candidato.get("classificacao") or "-"
            classificacao_def = candidato.get("classificacao_pcd") or "-"
            classificacao_nna = candidato.get("classificacao_nna") or "-"
            nome = candidato_obj.get("nome") or "-"
            rg = candidato_obj.get("rg") or "-"
            cpf = candidato_obj.get("cpf") or "-"
            agenda_data = agendas_por_candidato.get(str(candidato_uuid))
            if not agenda_data:
                cargo_codigo_candidato = candidato.get("codigo_cargo") or ""
                for agenda in agendas_map.values():
                    if agenda.get("cargo_codigo") == cargo_codigo_candidato:
                        agenda_data = agenda
                        break
            if not agenda_data:
                cargo_codigo_candidato = candidato.get("codigo_cargo") or ""
                cargo_descricao_candidato = (
                    candidato.get("descricao_cargo") or ""
                )
                if not cargo_descricao_candidato and cargo_codigo_candidato:
                    cargo_descricao_candidato = (
                        cargos_map.get(str(cargo_codigo_candidato))
                        or cargos_map.get(cargo_codigo_candidato)
                        or ""
                    )
                if not cargo_descricao_candidato and cargo_codigo_candidato:
                    cargo_descricao_candidato = (
                        f"Cargo {cargo_codigo_candidato}"
                    )
                elif not cargo_descricao_candidato:
                    cargo_descricao_candidato = "Cargo não informado"
                agenda_data = {
                    "uuid": None,
                    "cargo_uuid": None,
                    "cargo_nome": cargo_descricao_candidato,
                    "cargo_codigo": cargo_codigo_candidato,
                    "escolha_em": None,
                    "sessao": None,
                }
            cargo_codigo = agenda_data.get("cargo_codigo") or ""
            cargo_descricao = agenda_data.get("cargo_nome") or ""
            if not cargo_descricao and cargo_codigo:
                cargo_descricao = (
                    cargos_map.get(str(cargo_codigo))
                    or cargos_map.get(cargo_codigo)
                    or ""
                )
            if not cargo_descricao and cargo_codigo:
                cargo_descricao = f"Cargo {cargo_codigo}"
            elif not cargo_descricao:
                cargo_descricao = "Cargo não informado"
            situacao_escolha = escolha.get("situacao", "")
            if self.tipo == "RESULTADO_ESCOLHA":
                if situacao_escolha == "reconvocacao":
                    tipo_escolha_nome = "Reconvocação"
                    escolha_valor = "R"
                elif situacao_escolha == "nao-escolha":
                    tipo_escolha_nome = "Não Escolha"
                    escolha_valor = "N"
                elif situacao_escolha == "escolha":
                    tipo_escolha_nome = "Escolha"
                    escolha_valor = "S"
                else:
                    tipo_escolha_nome = "Outros"
                    escolha_valor = "-"
            elif self.tipo == "RESULTADO_ESCOLHA_RECONVOCACAO":
                tipo_escolha_nome = "Reconvocação"
                escolha_valor = "R"
            elif self.tipo == "RESULTADO_ESCOLHA_NAO":
                tipo_escolha_nome = "Não Escolha"
                escolha_valor = "N"
            elif self.tipo == "RESULTADO_ESCOLHA_SIM":
                tipo_escolha_nome = "Escolha"
                escolha_valor = "S"
            else:
                tipo_escolha_nome = "Outros"
                escolha_valor = "-"
            item_escolha = {
                "cargo_codigo": cargo_codigo,
                "cargo_descricao": cargo_descricao,
                "tipo_escolha": tipo_escolha_nome,
                "tipo_escolha_ordem": 1
                if tipo_escolha_nome == "Escolha"
                else 2
                if tipo_escolha_nome == "Não Escolha"
                else 3
                if tipo_escolha_nome == "Reconvocação"
                else 4,
                "agenda_uuid": agenda_data.get("uuid"),
                "agenda_nome": agenda_data.get("cargo_nome")
                or cargo_descricao,
                "agenda_data": agenda_data.get("escolha_em") or "-",
                "agenda_sessao": agenda_data.get("sessao") or "-",
                "classificacao_geral": classificacao_geral,
                "classificacao_def": classificacao_def,
                "classificacao_nna": classificacao_nna,
                "nome": nome,
                "rg": rg,
                "cpf": cpf,
                "escolha": escolha_valor,
            }
            if (
                self.tipo == "RESULTADO_ESCOLHA"
                and situacao_escolha == "escolha"
            ):
                vaga_escola = escolha.get("vaga_escola", {})
                escola = (
                    vaga_escola.get("escola", {})
                    if isinstance(vaga_escola, dict)
                    else {}
                )
                dre = escola.get("dre", {}) if isinstance(escola, dict) else {}
                item_escolha["dre_nome"] = (
                    dre.get("nome", "") if isinstance(dre, dict) else ""
                )
                item_escolha["dre_codigo"] = (
                    dre.get("codigo", "") if isinstance(dre, dict) else ""
                )
                item_escolha["escola_nome"] = (
                    escola.get("nome_oficial", "")
                    if isinstance(escola, dict)
                    else ""
                )
                item_escolha["escola_codigo_eol"] = (
                    escola.get("codigo_eol", "")
                    if isinstance(escola, dict)
                    else ""
                )
                item_escolha["tipo_ue"] = (
                    escola.get("tipo_ue", "")
                    if isinstance(escola, dict)
                    else ""
                )
                tipo_vaga_raw = escolha.get("tipo_vaga", "")
                item_escolha["tipo_vaga"] = (
                    "definitiva"
                    if tipo_vaga_raw == "definitiva"
                    else "precaria"
                    if tipo_vaga_raw == "precaria"
                    else ""
                )
            else:
                item_escolha["dre_nome"] = ""
                item_escolha["dre_codigo"] = ""
                item_escolha["escola_nome"] = ""
                item_escolha["escola_codigo_eol"] = ""
                item_escolha["tipo_ue"] = ""
                item_escolha["tipo_vaga"] = ""
            escolhas_com_candidatos.append(item_escolha)
        if self.tipo == "RESULTADO_ESCOLHA":
            cargos_list = self._agrupar_por_cargo_tipo_escolha_e_agenda(
                escolhas_com_candidatos
            )
            cargos_list = self._adicionar_resumo_dre_escola(
                cargos_list, escolhas_com_candidatos, processo_uuid
            )
        else:
            cargos_list = self._agrupar_por_cargo_e_agenda(
                escolhas_com_candidatos
            )
        logo_url = (
            request.build_absolute_uri(self.context.get("logo_url", ""))
            if self.context.get("logo_url")
            else ""
        )
        data_atual = timezone.now()
        self.context.update(
            {
                "cargos": cargos_list,
                "logo_url": logo_url,
                "data_atual": data_atual,
                "is_pdf": False,
            }
        )
        if formato == "xls" or formato == "csv":
            filename = f"resultado_escolha_{processo_uuid}.xlsx"
            logger.info("Gerando Excel: %s", filename)
            response = self.render_to_xls(
                cargos_list,
                self.context.get("cabecalho_padrao"),
                self.context.get("cabecalho"),
                filename=filename,
            )
            return (response, cargos_list)
        elif formato == "docx" or formato == "doc":
            filename = f"resultado_escolha_{processo_uuid}.docx"
            logger.info("Gerando Word: %s", filename)
            response = self.render_to_docx(
                cargos_list,
                self.context.get("cabecalho_padrao"),
                self.context.get("cabecalho"),
                self.context.get("texto_final"),
                filename=filename,
            )  # type: ignore[misc]
            return (response, cargos_list)
        elif formato == "pdf":
            filename = f"resultado_escolha_{processo_uuid}.pdf"
            logger.info("Gerando PDF: %s", filename)
            self.context["is_pdf"] = True
            response = self.render_to_pdf(
                self.TEMPLATE_NAME, self.context, filename=filename
            )
            return (response, cargos_list)
        else:
            logger.info("Gerando HTML")
            response = render(request, self.TEMPLATE_NAME, self.context)
            return (response, cargos_list)

    def _extrair_numero_sessao(self, sessao: str) -> str:
        """Extrai apenas o número da sessão, removendo a palavra "Sessão" se.

        Args:
            self: Instância do objeto.
            sessao: String com a sessão (ex: "Sessão 4", "4", etc.).

        Returns:
            Texto resultante da operação.

        Raises:
            Nenhuma exceção específica documentada.
        """
        if not sessao or sessao == "-":
            return "-"
        sessao_str = str(sessao).strip()
        sessao_limpa = re.sub(
            "^[Ss]ess[ãa]o\\s*", "", sessao_str, flags=re.IGNORECASE
        )
        sessao_limpa = sessao_limpa.strip()
        numeros = re.findall("\\d+", sessao_limpa)
        if numeros:
            return numeros[0]  # type: ignore[no-any-return]
        return sessao_limpa if sessao_limpa else "-"

    def _agrupar_por_cargo_e_agenda(self, escolhas: list) -> list:
        """Agrupa escolhas por cargo da agenda e depois por agenda.

        Args:
            self: Instância do objeto.
            escolhas: Lista de escolhas com suas informações.

        Returns:
            Lista com os registros resultantes.

        Raises:
            Nenhuma exceção específica documentada.
        """
        cargos_dict = {}  # type: ignore[var-annotated]
        for escolha in escolhas:
            cargo_codigo = escolha.get("cargo_codigo", "") or ""
            cargo_descricao = escolha.get("cargo_descricao", "") or ""
            agenda_uuid = escolha.get("agenda_uuid")
            agenda_nome = escolha.get("agenda_nome", "-")
            agenda_data = escolha.get("agenda_data", "-")
            agenda_sessao = escolha.get("agenda_sessao", "-")
            sessao_numero = self._extrair_numero_sessao(agenda_sessao)
            if not cargo_descricao:
                if cargo_codigo and cargo_codigo != "-":
                    cargo_descricao = f"Cargo {cargo_codigo}"
                else:
                    cargo_descricao = "Cargo não informado"
            if agenda_uuid:
                agenda_chave = str(agenda_uuid)
            else:
                agenda_chave = f"{agenda_nome}_{agenda_data}_{agenda_sessao}"
            if cargo_codigo not in cargos_dict:
                cargos_dict[cargo_codigo] = {
                    "codigo": cargo_codigo
                    if cargo_codigo and cargo_codigo != "-"
                    else "",
                    "descricao": cargo_descricao,
                    "agendas": {},
                }
            if agenda_chave not in cargos_dict[cargo_codigo]["agendas"]:
                cargos_dict[cargo_codigo]["agendas"][agenda_chave] = {
                    "uuid": agenda_uuid,
                    "nome": agenda_nome,
                    "data": agenda_data,
                    "sessao": sessao_numero,
                    "candidatos": [],
                }
            cargos_dict[cargo_codigo]["agendas"][agenda_chave][
                "candidatos"
            ].append(escolha)
        cargos_list = []
        for cargo_codigo, cargo_data in cargos_dict.items():
            agendas_list = []
            for agenda_chave, agenda_data in cargo_data["agendas"].items():
                agenda_data["candidatos"].sort(
                    key=lambda e: e["classificacao_geral"]
                    if isinstance(e["classificacao_geral"], int | float)
                    and e["classificacao_geral"] != "-"
                    else float("inf")
                )
                agendas_list.append(agenda_data)
            agendas_list.sort(
                key=lambda a: (
                    a["data"] if a["data"] != "-" else "",
                    int(a["sessao"])
                    if a["sessao"] != "-" and str(a["sessao"]).isdigit()
                    else float("inf")
                    if a["sessao"] != "-"
                    else "",
                )
            )
            cargos_list.append(
                {
                    "codigo": cargo_data["codigo"],
                    "descricao": cargo_data["descricao"],
                    "agendas": agendas_list,
                }
            )
        cargos_list.sort(key=lambda x: x["descricao"])
        return cargos_list

    def _agrupar_por_cargo_tipo_escolha_e_agenda(self, escolhas: list) -> list:
        """Agrupa escolhas por cargo da agenda, depois por tipo de escolha, e.

        Args:
            self: Instância do objeto.
            escolhas: Lista de escolhas com suas informações.

        Returns:
            Lista com os registros resultantes.

        Raises:
            Nenhuma exceção específica documentada.
        """
        cargos_dict = {}  # type: ignore[var-annotated]
        for escolha in escolhas:
            cargo_codigo = escolha.get("cargo_codigo", "") or ""
            cargo_descricao = escolha.get("cargo_descricao", "") or ""
            tipo_escolha = escolha.get("tipo_escolha", "Outros")
            tipo_escolha_ordem = escolha.get("tipo_escolha_ordem", 4)
            agenda_uuid = escolha.get("agenda_uuid")
            agenda_nome = escolha.get("agenda_nome", "-")
            agenda_data = escolha.get("agenda_data", "-")
            agenda_sessao = escolha.get("agenda_sessao", "-")
            sessao_numero = self._extrair_numero_sessao(agenda_sessao)
            if not cargo_descricao:
                if cargo_codigo and cargo_codigo != "-":
                    cargo_descricao = f"Cargo {cargo_codigo}"
                else:
                    cargo_descricao = "Cargo não informado"
            if agenda_uuid:
                agenda_chave = str(agenda_uuid)
            else:
                agenda_chave = f"{agenda_nome}_{agenda_data}_{agenda_sessao}"
            if cargo_codigo not in cargos_dict:
                cargos_dict[cargo_codigo] = {
                    "codigo": cargo_codigo
                    if cargo_codigo and cargo_codigo != "-"
                    else "",
                    "descricao": cargo_descricao,
                    "tipos_escolha": {},
                }
            if tipo_escolha not in cargos_dict[cargo_codigo]["tipos_escolha"]:
                cargos_dict[cargo_codigo]["tipos_escolha"][tipo_escolha] = {
                    "nome": tipo_escolha,
                    "ordem": tipo_escolha_ordem,
                    "agendas": {},
                }
            if (
                agenda_chave
                not in cargos_dict[cargo_codigo]["tipos_escolha"][
                    tipo_escolha
                ]["agendas"]
            ):
                cargos_dict[cargo_codigo]["tipos_escolha"][tipo_escolha][
                    "agendas"
                ][agenda_chave] = {
                    "uuid": agenda_uuid,
                    "nome": agenda_nome,
                    "data": agenda_data,
                    "sessao": sessao_numero,
                    "candidatos": [],
                }
            cargos_dict[cargo_codigo]["tipos_escolha"][tipo_escolha][
                "agendas"
            ][agenda_chave]["candidatos"].append(escolha)
        cargos_list = []
        for cargo_codigo, cargo_data in cargos_dict.items():
            tipos_escolha_list = []
            for _tipo_escolha_nome, tipo_escolha_data in cargo_data[
                "tipos_escolha"
            ].items():
                agendas_list = []
                for agenda_chave, agenda_data in tipo_escolha_data[
                    "agendas"
                ].items():
                    agenda_data["candidatos"].sort(
                        key=lambda e: e["classificacao_geral"]
                        if isinstance(e["classificacao_geral"], int | float)
                        and e["classificacao_geral"] != "-"
                        else float("inf")
                    )
                    agendas_list.append(agenda_data)
                agendas_list.sort(
                    key=lambda a: (
                        a["data"] if a["data"] != "-" else "",
                        int(a["sessao"])
                        if a["sessao"] != "-" and str(a["sessao"]).isdigit()
                        else float("inf")
                        if a["sessao"] != "-"
                        else "",
                    )
                )
                tipos_escolha_list.append(
                    {
                        "nome": tipo_escolha_data["nome"],
                        "ordem": tipo_escolha_data["ordem"],
                        "agendas": agendas_list,
                    }
                )
            tipos_escolha_list.sort(key=lambda t: t["ordem"])
            cargos_list.append(
                {
                    "codigo": cargo_data["codigo"],
                    "descricao": cargo_data["descricao"],
                    "tipos_escolha": tipos_escolha_list,
                }
            )
        cargos_list.sort(key=lambda x: x["descricao"])
        return cargos_list

    def _adicionar_resumo_dre_escola(
        self,
        cargos_list: list,
        escolhas_com_candidatos: list,
        processo_uuid: str,
    ) -> list:
        """Adiciona ao cargos_list o resumo DRE > ESCOLA com qtd de vagas e.

        Args:
            self: Instância do objeto.
            cargos_list: Parâmetro cargos list.
            escolhas_com_candidatos: Parâmetro escolhas com candidatos.
            processo_uuid: Parâmetro processo uuid.

        Returns:
            Lista com os registros resultantes.

        Raises:
            Nenhuma exceção específica documentada.
        """
        escolhas_realizadas = [
            e
            for e in escolhas_com_candidatos
            if e.get("tipo_escolha") == "Escolha"
            and (
                e.get("dre_nome")
                or e.get("escola_codigo_eol")
                or e.get("escola_nome")
            )
        ]
        if not escolhas_realizadas:
            return cargos_list
        por_cargo = {}  # type: ignore[var-annotated]
        for e in escolhas_realizadas:
            cargo_codigo = e.get("cargo_codigo", "") or ""
            dre_codigo = e.get("dre_codigo", "") or ""
            dre_nome = e.get("dre_nome", "") or "-"
            escola_eol = e.get("escola_codigo_eol", "") or ""
            escola_nome = e.get("escola_nome", "") or "-"
            tipo_ue = e.get("tipo_ue", "") or "-"
            tipo_vaga = e.get("tipo_vaga", "") or ""
            if not cargo_codigo:
                continue
            if cargo_codigo not in por_cargo:
                por_cargo[cargo_codigo] = {}
            if dre_codigo not in por_cargo[cargo_codigo]:
                por_cargo[cargo_codigo][dre_codigo] = {
                    "nome": dre_nome,
                    "escolas": {},
                }
            escola_chave = escola_eol or escola_nome or "-"
            if (
                escola_chave
                not in por_cargo[cargo_codigo][dre_codigo]["escolas"]
            ):
                por_cargo[cargo_codigo][dre_codigo]["escolas"][
                    escola_chave
                ] = {
                    "nome": escola_nome,
                    "codigo_eol": escola_eol,
                    "tipo_ue": tipo_ue,
                    "qtd_escolhas_definitivas": 0,
                    "qtd_escolhas_precarias": 0,
                }
            if tipo_vaga == "definitiva":
                por_cargo[cargo_codigo][dre_codigo]["escolas"][escola_chave][
                    "qtd_escolhas_definitivas"
                ] += 1
            elif tipo_vaga == "precaria":
                por_cargo[cargo_codigo][dre_codigo]["escolas"][escola_chave][
                    "qtd_escolhas_precarias"
                ] += 1
        mapa_vagas = {}
        try:
            vagas_response = self.escolhas_service.buscar_vagas_escolas(
                processo_uuid=str(processo_uuid) if processo_uuid else ""
            )
            vagas_list = vagas_response.json().get("vagas", [])
            for vaga in vagas_list:
                cargo_cod = vaga.get("cargo_codigo", "")
                escola = vaga.get("escola", {}) or {}
                dre = escola.get("dre", {}) or {}
                dre_cod = dre.get("codigo", "")
                codigo_eol = escola.get("codigo_eol", "")
                v_def = vaga.get("vagas_definitivas") or 0
                v_prec = vaga.get("vagas_precarias") or 0
                chave = (str(cargo_cod), str(dre_cod), str(codigo_eol))
                if chave not in mapa_vagas:
                    mapa_vagas[chave] = {"definitivas": 0, "precarias": 0}
                mapa_vagas[chave]["definitivas"] += v_def
                mapa_vagas[chave]["precarias"] += v_prec
        except Exception as exc:
            logger.warning(
                "Falha ao buscar vagas das escolas para resumo DRE/ESCOLA: %s. Exibindo apenas qtd de escolhas.",  # noqa: E501
                exc,
            )
        for cargo in cargos_list:
            cargo_codigo = cargo.get("codigo", "") or ""
            resumo_dres = []
            dres_data = por_cargo.get(cargo_codigo, {})
            for dre_codigo, dre_info in sorted(
                dres_data.items(), key=lambda x: (x[1]["nome"], x[0])
            ):
                escolas_list = []
                for escola_chave, esc_info in sorted(
                    dre_info["escolas"].items(),
                    key=lambda x: (x[1]["nome"], x[0]),
                ):
                    chave_vaga = (
                        cargo_codigo,
                        dre_codigo,
                        esc_info.get("codigo_eol", "") or "",
                    )
                    vagas = mapa_vagas.get(
                        chave_vaga, {"definitivas": 0, "precarias": 0}
                    )
                    escolas_list.append(
                        {
                            "nome": esc_info["nome"],
                            "codigo_eol": esc_info.get("codigo_eol", "")
                            or "-",
                            "tipo_ue": esc_info.get("tipo_ue", "") or "-",
                            "qtd_vagas_definitivas": vagas.get(
                                "definitivas", 0
                            ),
                            "qtd_vagas_precarias": vagas.get("precarias", 0),
                            "qtd_escolhas_definitivas": esc_info.get(
                                "qtd_escolhas_definitivas", 0
                            ),
                            "qtd_escolhas_precarias": esc_info.get(
                                "qtd_escolhas_precarias", 0
                            ),
                        }
                    )
                resumo_dres.append(
                    {"nome": dre_info["nome"], "escolas": escolas_list}
                )
            cargo["resumo_dre_escola"] = resumo_dres
        return cargos_list

    def render_to_xls(
        self,
        cargos_list: Any,
        cabecalho_padrao: Any,
        cabecalho: Any,
        filename: Any = "resultado_escolha.xlsx",
    ) -> Any:
        """Gera um arquivo Excel (XLSX) mantendo a estrutura hierárquica do.

        Args:
            self: Instância do objeto.
            cargos_list: Lista de cargos com suas agendas e candidatos.
            cabecalho_padrao: Cabeçalho padrão do relatório.
            cabecalho: Texto do cabeçalho do relatório.
            filename: Nome do arquivo Excel gerado.

        Returns:
            Resultado da operação.

        Raises:
            ImportError: Se ocorrer erro nesta operação.
        """
        if not OPENPYXL_AVAILABLE:
            raise ImportError(
                "openpyxl não está instalado. Instale com: pip install openpyxl>=3.1.0"  # noqa: E501
            )
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Resultado da Escolha"
            cargo_fill = PatternFill(
                start_color="667eea", end_color="667eea", fill_type="solid"
            )
            table_header_fill = PatternFill(
                start_color="ECF0F1", end_color="ECF0F1", fill_type="solid"
            )
            cargo_font = Font(bold=True, color="FFFFFF", size=12)
            header_font = Font(bold=True, size=10)
            normal_font = Font(size=10)
            title_font = Font(bold=True, size=14)
            border = Border(
                left=Side(style="thin"),
                right=Side(style="thin"),
                top=Side(style="thin"),
                bottom=Side(style="thin"),
            )
            center_align = Alignment(horizontal="center", vertical="center")
            center_wrap_align = Alignment(
                horizontal="center", vertical="center", wrap_text=True
            )
            left_align = Alignment(horizontal="left", vertical="center")
            row = 1
            temp_image_paths = []
            logo_url = self.context.get("logo_url", "")
            if self.context.get("usar_logotipo") and logo_url:
                image_path = None
                try:
                    if logo_url.startswith("http://") or logo_url.startswith(
                        "https://"
                    ):
                        with tempfile.NamedTemporaryFile(
                            suffix=".png", delete=False
                        ) as tmpf:
                            resp = requests.get(logo_url, timeout=15)
                            resp.raise_for_status()
                            tmpf.write(resp.content)
                            image_path = tmpf.name
                            temp_image_paths.append(image_path)
                    elif os.path.exists(logo_url):
                        image_path = logo_url
                    if image_path:
                        img = XLImage(image_path)
                        try:
                            img.width = 220
                            img.height = 90
                        except Exception:
                            pass
                        ws.add_image(img, "B1")
                        row = max(row, 8)
                except Exception as exc:
                    logger.warning(
                        "Não foi possível inserir o logotipo no XLS (resultado_escolha): %s",  # noqa: E501
                        exc,
                    )
            if cabecalho_padrao:
                ws.merge_cells(f"A{row}:H{row}")
                cell = ws[f"A{row}"]
                cabecalho_padrao_texto = self.processar_cabecalho_html(
                    cabecalho_padrao
                )
                cell.value = cabecalho_padrao_texto
                cell.font = title_font
                cell.alignment = center_wrap_align
                row += 2
            if cabecalho:
                ws.merge_cells(f"A{row}:H{row}")
                cell = ws[f"A{row}"]
                cabecalho_texto = self.processar_cabecalho_html(cabecalho)
                cell.value = cabecalho_texto
                cell.font = title_font
                cell.alignment = center_wrap_align
                row += 2
            ws.merge_cells(f"A{row}:H{row}")
            cell = ws[f"A{row}"]
            if self.tipo == "RESULTADO_ESCOLHA":
                cell.value = "RESULTADO DE ESCOLHA DE VAGAS"
            else:
                cell.value = "RESULTADO DA ESCOLHA DE VAGAS - GERAL"
            cell.font = Font(bold=True, size=16)
            cell.alignment = center_align
            row += 1
            ws.merge_cells(f"A{row}:H{row}")
            cell = ws[f"A{row}"]
            data_atual = timezone.now()
            meses_pt = {
                1: "janeiro",
                2: "fevereiro",
                3: "março",
                4: "abril",
                5: "maio",
                6: "junho",
                7: "julho",
                8: "agosto",
                9: "setembro",
                10: "outubro",
                11: "novembro",
                12: "dezembro",
            }
            mes_nome = meses_pt.get(data_atual.month, "")
            data_formatada = (
                f"{data_atual.day} de {mes_nome} de {data_atual.year}"
            )
            cell.value = f"REALIZADO EM: {data_formatada}"
            cell.font = Font(size=12)
            cell.alignment = center_align
            row += 2
            for cargo in cargos_list:
                cargo_descricao = cargo.get("descricao", "")
                ws.merge_cells(f"A{row}:H{row}")
                cell = ws[f"A{row}"]
                cell.value = cargo_descricao
                cell.font = cargo_font
                cell.fill = cargo_fill
                cell.alignment = left_align
                row += 1
                if "tipos_escolha" in cargo:
                    for tipo_escolha in cargo.get("tipos_escolha", []):
                        tipo_nome = tipo_escolha.get("nome", "")
                        ws.merge_cells(f"A{row}:H{row}")
                        cell = ws[f"A{row}"]
                        cell.value = f"  {tipo_nome}"
                        cell.font = Font(bold=True, size=11)
                        cell.fill = PatternFill(
                            start_color="D5D8DC",
                            end_color="D5D8DC",
                            fill_type="solid",
                        )
                        cell.alignment = left_align
                        row += 1
                        headers = [
                            "Bloco",
                            "Geral",
                            "NNA",
                            "Def.",
                            "NOME",
                            "R.G.",
                            "CPF",
                            "ESCOLHA",
                        ]
                        for col, header in enumerate(headers, start=1):
                            cell = ws.cell(row=row, column=col)
                            cell.value = header
                            cell.fill = table_header_fill
                            cell.font = header_font
                            cell.alignment = center_align
                            cell.border = border
                        row += 1
                        for agenda in tipo_escolha.get("agendas", []):
                            sessao = agenda.get("sessao", "-")
                            for candidato in agenda.get("candidatos", []):
                                ws.cell(row=row, column=1).value = sessao
                                ws.cell(
                                    row=row, column=2
                                ).value = candidato.get(
                                    "classificacao_geral", "-"
                                )
                                ws.cell(
                                    row=row, column=3
                                ).value = candidato.get(
                                    "classificacao_nna", "-"
                                )
                                ws.cell(
                                    row=row, column=4
                                ).value = candidato.get(
                                    "classificacao_def", "-"
                                )
                                ws.cell(
                                    row=row, column=5
                                ).value = candidato.get("nome", "-")
                                ws.cell(
                                    row=row, column=6
                                ).value = candidato.get("rg", "-")
                                ws.cell(
                                    row=row, column=7
                                ).value = candidato.get("cpf", "-")
                                ws.cell(
                                    row=row, column=8
                                ).value = candidato.get("escolha", "-")
                                for col in range(1, 9):
                                    cell = ws.cell(row=row, column=col)
                                    cell.border = border
                                    cell.font = normal_font
                                    if col in [1, 2, 3, 4, 8]:
                                        cell.alignment = center_align
                                    else:
                                        cell.alignment = left_align
                                row += 1
                else:
                    headers = [
                        "Bloco",
                        "Geral",
                        "NNA",
                        "Def.",
                        "NOME",
                        "R.G.",
                        "CPF",
                        "ESCOLHA",
                    ]
                    for col, header in enumerate(headers, start=1):
                        cell = ws.cell(row=row, column=col)
                        cell.value = header
                        cell.fill = table_header_fill
                        cell.font = header_font
                        cell.alignment = center_align
                        cell.border = border
                    row += 1
                    for agenda in cargo.get("agendas", []):
                        sessao = agenda.get("sessao", "-")
                        for candidato in agenda.get("candidatos", []):
                            ws.cell(row=row, column=1).value = sessao
                            ws.cell(row=row, column=2).value = candidato.get(
                                "classificacao_geral", "-"
                            )
                            ws.cell(row=row, column=3).value = candidato.get(
                                "classificacao_nna", "-"
                            )
                            ws.cell(row=row, column=4).value = candidato.get(
                                "classificacao_def", "-"
                            )
                            ws.cell(row=row, column=5).value = candidato.get(
                                "nome", "-"
                            )
                            ws.cell(row=row, column=6).value = candidato.get(
                                "rg", "-"
                            )
                            ws.cell(row=row, column=7).value = candidato.get(
                                "cpf", "-"
                            )
                            ws.cell(row=row, column=8).value = candidato.get(
                                "escolha", "-"
                            )
                            for col in range(1, 9):
                                cell = ws.cell(row=row, column=col)
                                cell.border = border
                                cell.font = normal_font
                                if col in [1, 2, 3, 4, 8]:
                                    cell.alignment = center_align
                                else:
                                    cell.alignment = left_align
                            row += 1
                if cargo.get("resumo_dre_escola"):
                    ws.merge_cells(f"A{row}:G{row}")
                    cell = ws[f"A{row}"]
                    cell.value = "  Resumo por DRE e Escola (vagas e escolhas realizadas)"  # noqa: E501
                    cell.font = Font(bold=True, size=11)
                    cell.fill = PatternFill(
                        start_color="95a5a6",
                        end_color="95a5a6",
                        fill_type="solid",
                    )
                    cell.alignment = left_align
                    row += 1
                    headers_resumo = [
                        "Escola",
                        "Tipo UE",
                        "Código EOL",
                        "Vagas Definitivas",
                        "Escolhas Definitivas",
                        "Vagas Precárias",
                        "Escolhas Precárias",
                    ]
                    for dre in cargo.get("resumo_dre_escola", []):
                        ws.merge_cells(f"A{row}:G{row}")
                        cell = ws[f"A{row}"]
                        cell.value = f'  DRE - {dre.get('nome', '')}'
                        cell.font = Font(bold=True, size=10)
                        cell.fill = PatternFill(
                            start_color="BDC3C7",
                            end_color="BDC3C7",
                            fill_type="solid",
                        )
                        cell.alignment = left_align
                        row += 1
                        for col, header in enumerate(headers_resumo, start=1):
                            cell = ws.cell(row=row, column=col)
                            cell.value = header
                            cell.fill = table_header_fill
                            cell.font = header_font
                            cell.alignment = center_align
                            cell.border = border
                        row += 1
                        for escola in dre.get("escolas", []):
                            ws.cell(row=row, column=1).value = escola.get(
                                "nome", "-"
                            )
                            ws.cell(row=row, column=2).value = escola.get(
                                "tipo_ue", "-"
                            )
                            ws.cell(row=row, column=3).value = escola.get(
                                "codigo_eol", "-"
                            )
                            ws.cell(row=row, column=4).value = escola.get(
                                "qtd_vagas_definitivas", 0
                            )
                            ws.cell(row=row, column=5).value = escola.get(
                                "qtd_escolhas_definitivas", 0
                            )
                            ws.cell(row=row, column=6).value = escola.get(
                                "qtd_vagas_precarias", 0
                            )
                            ws.cell(row=row, column=7).value = escola.get(
                                "qtd_escolhas_precarias", 0
                            )
                            for col in range(1, 8):
                                cell = ws.cell(row=row, column=col)
                                cell.border = border
                                cell.font = normal_font
                                cell.alignment = (
                                    center_align
                                    if col in [2, 3, 4, 5, 6, 7]
                                    else left_align
                                )
                            row += 1
                    row += 1
                row += 1
            max_escola_length = 0
            for cargo in cargos_list:
                if cargo.get("resumo_dre_escola"):
                    for dre in cargo.get("resumo_dre_escola", []):
                        for escola in dre.get("escolas", []):
                            escola_nome = str(escola.get("nome", "-"))
                            max_escola_length = max(
                                max_escola_length, len(escola_nome)
                            )
            escola_width = 40
            if max_escola_length > 0:
                escola_width = min(max(max_escola_length + 5, 40), 50)
            column_widths = {
                "A": max(10, escola_width),
                "B": 12,
                "C": 12,
                "D": 14,
                "E": 40,
                "F": 18,
                "G": 16,
                "H": 10,
            }
            for col_letter, width in column_widths.items():
                ws.column_dimensions[col_letter].width = width
            if max_escola_length > 0:
                resumo_numeric_widths = {"D": 15, "E": 15, "F": 18, "G": 18}
                for col_letter, width in resumo_numeric_widths.items():
                    if col_letter == "E":
                        current_width = (
                            ws.column_dimensions[col_letter].width or 40
                        )
                        ws.column_dimensions[col_letter].width = max(
                            current_width, 40
                        )
                    else:
                        current_width = (
                            ws.column_dimensions[col_letter].width or width
                        )
                        ws.column_dimensions[col_letter].width = max(
                            current_width, width
                        )
            texto_final = self.context.get("texto_final")
            if texto_final:
                row += 1
                ws.merge_cells(f"A{row}:H{row}")
                cell = ws[f"A{row}"]
                cell.value = self.processar_cabecalho_html(texto_final)
                cell.font = normal_font
                cell.alignment = Alignment(
                    horizontal="left", vertical="top", wrap_text=True
                )
            buffer = BytesIO()
            wb.save(buffer)
            buffer.seek(0)
            for p in temp_image_paths:
                try:
                    if os.path.exists(p):
                        os.unlink(p)
                except Exception:
                    pass
            response = HttpResponse(
                buffer.read(),
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            response["Content-Disposition"] = (
                f'attachment; filename="{filename}"'
            )
            return response
        except Exception as exc:
            logger.error("Erro ao gerar Excel: %s", exc, exc_info=True)
            raise

    def render_to_docx(
        self,
        cargos_list: Any,
        cabecalho: Any,
        texto_final: Any = None,
        filename: Any = "resultado_escolha.docx",
    ) -> Any:
        """Gera um arquivo Word (DOCX) mantendo a estrutura hierárquica do.

        Args:
            self: Instância do objeto.
            cargos_list: Lista de cargos com suas agendas e candidatos.
            cabecalho: Texto do cabeçalho do relatório.
            texto_final: Texto final do relatório (opcional).
            filename: Nome do arquivo Word gerado.

        Returns:
            Resultado da operação.

        Raises:
            ImportError: Se ocorrer erro nesta operação.
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx não está instalado. Instale com: pip install python-docx>=1.1.0"  # noqa: E501
            )
        try:
            doc = Document()
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            RGBColor(102, 126, 234)
            RGBColor(236, 240, 241)
            if cabecalho:
                cabecalho_texto = self.processar_cabecalho_html(cabecalho)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cabecalho_texto)
                run.font.size = Pt(14)
                run.font.bold = True
                doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if self.tipo == "RESULTADO_ESCOLHA":
                run = p.add_run("RESULTADO DE ESCOLHA DE VAGAS")
            else:
                run = p.add_run("RESULTADO DA ESCOLHA DE VAGAS - GERAL")
            run.font.size = Pt(16)
            run.font.bold = True
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            data_atual = timezone.now()
            meses_pt = {
                1: "janeiro",
                2: "fevereiro",
                3: "março",
                4: "abril",
                5: "maio",
                6: "junho",
                7: "julho",
                8: "agosto",
                9: "setembro",
                10: "outubro",
                11: "novembro",
                12: "dezembro",
            }
            mes_nome = meses_pt.get(data_atual.month, "")
            data_formatada = (
                f"{data_atual.day} de {mes_nome} de {data_atual.year}"
            )
            run = p.add_run(f"REALIZADO EM: {data_formatada}")
            run.font.size = Pt(12)
            doc.add_paragraph()
            doc.add_paragraph()
            for cargo in cargos_list:
                cargo_descricao = cargo.get("descricao", "")
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(cargo_descricao)
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                p_pr = p._element.get_or_add_pPr()
                existing_shd = p_pr.find(qn("w:shd"))
                if existing_shd is not None:
                    p_pr.remove(existing_shd)
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:fill"), "667eea")
                shading_elm.set(qn("w:val"), "clear")
                p_pr.append(shading_elm)
                if "tipos_escolha" in cargo:
                    for tipo_escolha in cargo.get("tipos_escolha", []):
                        tipo_nome = tipo_escolha.get("nome", "")
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = p.add_run(f"  {tipo_nome}")
                        run.font.size = Pt(11)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        p_pr = p._element.get_or_add_pPr()
                        existing_shd = p_pr.find(qn("w:shd"))
                        if existing_shd is not None:
                            p_pr.remove(existing_shd)
                        shading_elm = OxmlElement("w:shd")
                        shading_elm.set(qn("w:fill"), "D5D8DC")
                        shading_elm.set(qn("w:val"), "clear")
                        p_pr.append(shading_elm)
                        headers = [
                            "Bloco",
                            "Geral",
                            "NNA",
                            "Def.",
                            "NOME",
                            "R.G.",
                            "CPF",
                            "ESCOLHA",
                        ]
                        table = doc.add_table(rows=1, cols=len(headers))
                        table.style = "Light Grid Accent 1"
                        header_cells = table.rows[0].cells
                        for i, header in enumerate(headers):
                            cell = header_cells[i]
                            cell.text = header
                            cell.paragraphs[0].alignment = (
                                WD_ALIGN_PARAGRAPH.CENTER
                                if i in [0, 1, 2, 3, 7]
                                else WD_ALIGN_PARAGRAPH.LEFT
                            )
                            cell.paragraphs[0].runs[0].font.bold = True
                            cell.paragraphs[0].runs[0].font.size = Pt(10)
                            tc_pr = cell._element.get_or_add_tcPr()
                            existing_shd = tc_pr.find(qn("w:shd"))
                            if existing_shd is not None:
                                tc_pr.remove(existing_shd)
                            shading_elm = OxmlElement("w:shd")
                            shading_elm.set(qn("w:fill"), "ECF0F1")
                            shading_elm.set(qn("w:val"), "clear")
                            tc_pr.append(shading_elm)
                        for agenda in tipo_escolha.get("agendas", []):
                            sessao = agenda.get("sessao", "-")
                            for candidato in agenda.get("candidatos", []):
                                row_cells = table.add_row().cells
                                row_cells[0].text = str(sessao)
                                row_cells[1].text = str(
                                    candidato.get("classificacao_geral", "-")
                                )
                                row_cells[2].text = str(
                                    candidato.get("classificacao_nna", "-")
                                )
                                row_cells[3].text = str(
                                    candidato.get("classificacao_def", "-")
                                )
                                row_cells[4].text = str(
                                    candidato.get("nome", "-")
                                )
                                row_cells[5].text = str(
                                    candidato.get("rg", "-")
                                )
                                row_cells[6].text = str(
                                    candidato.get("cpf", "-")
                                )
                                row_cells[7].text = str(
                                    candidato.get("escolha", "-")
                                )
                                for i, cell in enumerate(row_cells):
                                    if i in [0, 1, 2, 3, 7]:
                                        cell.paragraphs[
                                            0
                                        ].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    else:
                                        cell.paragraphs[
                                            0
                                        ].alignment = WD_ALIGN_PARAGRAPH.LEFT
                                    cell.paragraphs[0].runs[0].font.size = Pt(
                                        10
                                    )
                else:
                    headers = [
                        "Bloco",
                        "Geral",
                        "NNA",
                        "Def.",
                        "NOME",
                        "R.G.",
                        "CPF",
                        "ESCOLHA",
                    ]
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = "Light Grid Accent 1"
                    header_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        cell = header_cells[i]
                        cell.text = header
                        cell.paragraphs[0].alignment = (
                            WD_ALIGN_PARAGRAPH.CENTER
                            if i in [0, 1, 2, 3, 7]
                            else WD_ALIGN_PARAGRAPH.LEFT
                        )
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].runs[0].font.size = Pt(10)
                        tc_pr = cell._element.get_or_add_tcPr()
                        existing_shd = tc_pr.find(qn("w:shd"))
                        if existing_shd is not None:
                            tc_pr.remove(existing_shd)
                        shading_elm = OxmlElement("w:shd")
                        shading_elm.set(qn("w:fill"), "ECF0F1")
                        shading_elm.set(qn("w:val"), "clear")
                        tc_pr.append(shading_elm)
                    for agenda in cargo.get("agendas", []):
                        sessao = agenda.get("sessao", "-")
                        for candidato in agenda.get("candidatos", []):
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(sessao)
                            row_cells[1].text = str(
                                candidato.get("classificacao_geral", "-")
                            )
                            row_cells[2].text = str(
                                candidato.get("classificacao_nna", "-")
                            )
                            row_cells[3].text = str(
                                candidato.get("classificacao_def", "-")
                            )
                            row_cells[4].text = str(candidato.get("nome", "-"))
                            row_cells[5].text = str(candidato.get("rg", "-"))
                            row_cells[6].text = str(candidato.get("cpf", "-"))
                            row_cells[7].text = str(
                                candidato.get("escolha", "-")
                            )
                            for i, cell in enumerate(row_cells):
                                if i in [0, 1, 2, 3, 7]:
                                    cell.paragraphs[
                                        0
                                    ].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                else:
                                    cell.paragraphs[
                                        0
                                    ].alignment = WD_ALIGN_PARAGRAPH.LEFT
                                cell.paragraphs[0].runs[0].font.size = Pt(10)
                if cargo.get("resumo_dre_escola"):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(
                        "  Resumo por DRE e Escola (vagas e escolhas realizadas)"  # noqa: E501
                    )
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    p_pr = p._element.get_or_add_pPr()
                    existing_shd = p_pr.find(qn("w:shd"))
                    if existing_shd is not None:
                        p_pr.remove(existing_shd)
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:fill"), "95a5a6")
                    shading_elm.set(qn("w:val"), "clear")
                    p_pr.append(shading_elm)
                    for dre in cargo.get("resumo_dre_escola", []):
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        run = p.add_run(f'  DRE - {dre.get('nome', '')}')
                        run.font.size = Pt(10)
                        run.font.bold = True
                        p_pr = p._element.get_or_add_pPr()
                        existing_shd = p_pr.find(qn("w:shd"))
                        if existing_shd is not None:
                            p_pr.remove(existing_shd)
                        shading_elm = OxmlElement("w:shd")
                        shading_elm.set(qn("w:fill"), "BDC3C7")
                        shading_elm.set(qn("w:val"), "clear")
                        p_pr.append(shading_elm)
                        headers_resumo = [
                            "Escola",
                            "Tipo UE",
                            "Código EOL",
                            "Vagas Definitivas",
                            "Escolhas Definitivas",
                            "Vagas Precárias",
                            "Escolhas Precárias",
                        ]
                        table = doc.add_table(rows=1, cols=7)
                        table.style = "Light Grid Accent 1"
                        header_cells = table.rows[0].cells
                        for i, header in enumerate(headers_resumo):
                            cell = header_cells[i]
                            cell.text = header
                            cell.paragraphs[0].alignment = (
                                WD_ALIGN_PARAGRAPH.CENTER
                                if i in [1, 2, 3, 4, 5, 6]
                                else WD_ALIGN_PARAGRAPH.LEFT
                            )
                            cell.paragraphs[0].runs[0].font.bold = True
                            cell.paragraphs[0].runs[0].font.size = Pt(10)
                            tc_pr = cell._element.get_or_add_tcPr()
                            existing_shd = tc_pr.find(qn("w:shd"))
                            if existing_shd is not None:
                                tc_pr.remove(existing_shd)
                            shading_elm = OxmlElement("w:shd")
                            shading_elm.set(qn("w:fill"), "ECF0F1")
                            shading_elm.set(qn("w:val"), "clear")
                            tc_pr.append(shading_elm)
                        for escola in dre.get("escolas", []):
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(escola.get("nome", "-"))
                            row_cells[1].text = str(escola.get("tipo_ue", "-"))
                            row_cells[2].text = str(
                                escola.get("codigo_eol", "-")
                            )
                            row_cells[3].text = str(
                                escola.get("qtd_vagas_definitivas", 0)
                            )
                            row_cells[4].text = str(
                                escola.get("qtd_escolhas_definitivas", 0)
                            )
                            row_cells[5].text = str(
                                escola.get("qtd_vagas_precarias", 0)
                            )
                            row_cells[6].text = str(
                                escola.get("qtd_escolhas_precarias", 0)
                            )
                            for i, cell in enumerate(row_cells):
                                cell.paragraphs[0].alignment = (
                                    WD_ALIGN_PARAGRAPH.CENTER
                                    if i in [1, 2, 3, 4, 5, 6]
                                    else WD_ALIGN_PARAGRAPH.LEFT
                                )
                                cell.paragraphs[0].runs[0].font.size = Pt(10)
                    doc.add_paragraph()
                doc.add_paragraph()
            if texto_final:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(self.processar_cabecalho_html(texto_final))
                run.font.size = Pt(10)
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            response = HttpResponse(
                buffer.read(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = (
                f'attachment; filename="{filename}"'
            )
            return response
        except Exception as exc:
            logger.error("Erro ao gerar Word: %s", exc, exc_info=True)
            raise

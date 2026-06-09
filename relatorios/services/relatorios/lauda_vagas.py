"""Implementação concreta do relatório de Lauda de Vagas."""

from __future__ import annotations

import logging
import os
import tempfile
from io import BytesIO
from typing import Any

import requests
from django.conf import settings
from django.http import HttpResponse
from django.shortcuts import render

from relatorios.services.base.relatorio_base import RelatorioBase
from relatorios.services.escolhas_api_service import EscolhasService
from relatorios.utils import convert_uuids_to_strings

try:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
logger = logging.getLogger(__name__)


class LaudaVagas(RelatorioBase):
    """Classe concreta responsável por gerar o relatório de Lauda de Vagas."""

    TEMPLATE_NAME = "relatorios/vagas_escolas.html"

    def __init__(self, **kwargs: Any) -> None:
        """Inicializa o service com as dependências necessárias.

        Args:
            self: Instância do objeto.
            **kwargs: Argumentos nomeados variáveis.

        Raises:
            Nenhuma exceção específica documentada.
        """
        super().__init__(**kwargs)
        self.escolhas_service = EscolhasService(
            base_url=settings.ESCOLHAS_API_URL
        )

    def gerar(
        self,
        processo_uuid: str,
        request: Any,
        formato: str = "html",
        cabecalho: str = "",
        **kwargs: Any,
    ) -> Any:
        """Gera o relatório de Lauda de Vagas.

        Args:
            self: Instância do objeto.
            processo_uuid: UUID do processo de convocação.
            request: Objeto request do Django.
            formato: Formato do relatório ('html', 'pdf', 'xls' ou 'docx').
            cabecalho: Texto do cabeçalho do relatório (opcional).
            **kwargs: Argumentos nomeados variáveis.

        Returns:
            Resultado da operação.

        Raises:
            Nenhuma exceção específica documentada.
        """
        try:
            vagas_escolas = self.escolhas_service.buscar_vagas_escolas(
                processo_uuid=str(processo_uuid) if processo_uuid else ""
            )
        except Exception as exc:
            logger.error(
                "Falha ao buscar vagas de escolas da API externa: %s", exc
            )
            raise
        vagas = vagas_escolas.json().get("vagas", [])
        vagas_agrupadas = self._agrupar_vagas(vagas)
        cargos_list = self._preparar_dados_template(vagas_agrupadas)
        cargos_list = convert_uuids_to_strings(cargos_list)
        logo_url = (
            request.build_absolute_uri(self.context.get("logo_url", ""))
            if self.context.get("logo_url")
            else ""
        )
        self.context.update({"logo_url": logo_url, "cargos": cargos_list})
        if formato == "xls" or formato == "csv":
            filename = f"relatorio_vagas_{processo_uuid}.xlsx"
            logger.info("Gerando Excel: %s", filename)
            response = self.render_to_xls(
                context=self.context, filename=filename
            )
            return (response, cargos_list)
        elif formato == "docx" or formato == "doc":
            filename = f"relatorio_vagas_{processo_uuid}.docx"
            logger.info("Gerando Word: %s", filename)
            response = self.render_to_docx(
                cargos_list,
                self.context,
                self.context["texto_final"],
                filename=filename,
            )
            return (response, cargos_list)
        elif formato == "pdf":
            filename = f"relatorio_vagas_{processo_uuid}.pdf"
            logger.info("Gerando PDF: %s", filename)
            self.context.update({"is_pdf": True, "cargos": cargos_list})
            response = self.render_to_pdf(
                self.TEMPLATE_NAME, self.context, filename=filename
            )
            return (response, cargos_list)
        else:
            logger.info("Gerando HTML")
            self.context["cargos"] = cargos_list
            response = render(request, self.TEMPLATE_NAME, self.context)
            return (response, cargos_list)

    def _agrupar_vagas(self, vagas: list) -> dict:
        """Agrupa vagas por cargo_codigo e depois por DRE codigo.

        Args:
            self: Instância do objeto.
            vagas: Lista de vagas.

        Returns:
            Dicionário com os dados processados.

        Raises:
            Nenhuma exceção específica documentada.
        """
        vagas_agrupadas = {}  # type: ignore[var-annotated]
        for vaga in vagas:
            cargo_codigo = vaga.get("cargo_codigo")
            dre_codigo = vaga.get("escola", {}).get("dre", {}).get("codigo")
            if cargo_codigo not in vagas_agrupadas:
                vagas_agrupadas[cargo_codigo] = {}
            if dre_codigo not in vagas_agrupadas[cargo_codigo]:
                vagas_agrupadas[cargo_codigo][dre_codigo] = []
            vagas_agrupadas[cargo_codigo][dre_codigo].append(vaga)
        return vagas_agrupadas

    def _preparar_dados_template(self, vagas_agrupadas: dict) -> list:
        """Prepara a estrutura de dados para o template.

        Args:
            self: Instância do objeto.
            vagas_agrupadas: Dicionário com vagas agrupadas por cargo e DRE.

        Returns:
            Lista com os registros resultantes.

        Raises:
            Nenhuma exceção específica documentada.
        """
        cargos_list = []
        for cargo_codigo, dres in vagas_agrupadas.items():
            primeira_vaga = None
            for dre_codigo, vagas_list in dres.items():
                if vagas_list:
                    primeira_vaga = vagas_list[0]
                    break
            if primeira_vaga:
                dres_list = []
                for dre_codigo, vagas_list in dres.items():
                    if vagas_list:
                        primeira_vaga_dre = vagas_list[0]
                        dres_list.append(
                            {
                                "codigo": dre_codigo,
                                "nome": primeira_vaga_dre.get("escola", {})
                                .get("dre", {})
                                .get("nome", ""),
                                "vagas": vagas_list,
                            }
                        )
                cargos_list.append(
                    {
                        "codigo": cargo_codigo,
                        "descricao": primeira_vaga.get("cargo_descricao", ""),
                        "dres": dres_list,
                    }
                )
        return cargos_list

    def render_to_xls(
        self, context: Any = None, filename: Any = "relatorio.xlsx"
    ) -> Any:
        """Gera um arquivo Excel (XLSX) mantendo a estrutura hierárquica do.

        Args:
            self: Instância do objeto.
            context: Dicionário com o contexto do template.
            filename: Nome do arquivo Excel gerado.

        Returns:
            Resultado da operação.

        Raises:
            ImportError: Se ocorrer erro nesta operação.
        """
        if context is None:
            context = {}
        if not OPENPYXL_AVAILABLE:
            raise ImportError(
                "openpyxl não está instalado. Instale com: pip install openpyxl>=3.1.0"  # noqa: E501
            )
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Relatório de Vagas"
            cargo_fill = PatternFill(
                start_color="667eea", end_color="667eea", fill_type="solid"
            )
            dre_fill = PatternFill(
                start_color="34495e", end_color="34495e", fill_type="solid"
            )
            table_header_fill = PatternFill(
                start_color="ECF0F1", end_color="ECF0F1", fill_type="solid"
            )
            cargo_font = Font(bold=True, color="FFFFFF", size=12)
            dre_font = Font(bold=True, color="FFFFFF", size=11)
            header_font = Font(bold=True, size=10)
            normal_font = Font(size=10)
            title_font = Font(bold=True, size=14)
            border = Border(
                left=Side(style="thin"),
                right=Side(style="thin"),
                top=Side(style="thin"),
                bottom=Side(style="thin"),
            )
            center_align = Alignment(horizontal="center", vertical="center")
            center_wrap_align = Alignment(
                horizontal="center", vertical="center", wrap_text=True
            )
            left_align = Alignment(horizontal="left", vertical="center")
            row = 1
            temp_image_paths = []
            logo_url = (
                (context or self.context).get("logo_url")
                if context or self.context
                else ""
            )
            if context.get("usar_logotipo") and logo_url:
                image_path = None
                try:
                    if logo_url.startswith("http://") or logo_url.startswith(
                        "https://"
                    ):
                        with tempfile.NamedTemporaryFile(
                            suffix=".png", delete=False
                        ) as tmpf:
                            resp = requests.get(logo_url, timeout=15)
                            resp.raise_for_status()
                            tmpf.write(resp.content)
                            image_path = tmpf.name
                            temp_image_paths.append(image_path)
                    elif os.path.exists(logo_url):
                        image_path = logo_url
                    if image_path:
                        img = XLImage(image_path)
                        try:
                            img.width = 220
                            img.height = 90
                        except Exception:
                            pass
                        ws.add_image(img, "B1")
                        row = max(row, 8)
                except Exception as exc:
                    logger.warning(
                        "Não foi possível inserir o logotipo no XLS: %s", exc
                    )
            cabecalho_padrao = self.context.get("cabecalho_padrao", "")
            if cabecalho_padrao:
                ws.merge_cells(f"A{row}:D{row}")
                cell = ws[f"A{row}"]
                cell.value = self.processar_cabecalho_html(cabecalho_padrao)
                cell.font = title_font
                cell.alignment = center_wrap_align
                row += 2
            if self.context.get("cabecalho"):
                ws.merge_cells(f"A{row}:D{row}")
                cell = ws[f"A{row}"]
                cell.value = self.processar_cabecalho_html(
                    self.context["cabecalho"]
                )
                cell.font = title_font
                cell.alignment = center_wrap_align
                row += 2
            for cargo in self.context["cargos"]:
                cargo_descricao = cargo.get("descricao", "")
                ws.merge_cells(f"A{row}:D{row}")
                cell = ws[f"A{row}"]
                cell.value = f"Cargo: {cargo_descricao}"
                cell.font = cargo_font
                cell.fill = cargo_fill
                cell.alignment = left_align
                row += 1
                for dre in cargo.get("dres", []):
                    dre_nome = dre.get("nome", "")
                    ws.merge_cells(f"A{row}:D{row}")
                    cell = ws[f"A{row}"]
                    cell.value = f"DRE - {dre_nome}"
                    cell.font = dre_font
                    cell.fill = dre_fill
                    cell.alignment = left_align
                    row += 1
                    headers = [
                        "Tipo de unidade",
                        "Unidade",
                        "Vagas Definitivas",
                        "Vagas Precárias",
                    ]
                    for col, header in enumerate(headers, start=1):
                        cell = ws.cell(row=row, column=col)
                        cell.value = header
                        cell.fill = table_header_fill
                        cell.font = header_font
                        cell.alignment = center_align
                        cell.border = border
                    row += 1
                    for vaga in dre.get("vagas", []):
                        escola = vaga.get("escola", {})
                        ws.cell(row=row, column=1).value = escola.get(
                            "tipo_ue", "-"
                        )
                        ws.cell(row=row, column=2).value = escola.get(
                            "nome_oficial", "-"
                        )
                        ws.cell(row=row, column=3).value = vaga.get(
                            "vagas_definitivas", 0
                        )
                        ws.cell(row=row, column=4).value = vaga.get(
                            "vagas_precarias", 0
                        )
                        for col in range(1, 5):
                            cell = ws.cell(row=row, column=col)
                            cell.border = border
                            cell.font = normal_font
                            if col in [3, 4]:
                                cell.alignment = center_align
                            else:
                                cell.alignment = left_align
                        row += 1
                    row += 1
                row += 1
            column_widths = {"A": 20, "B": 60, "C": 20, "D": 20}
            for col_letter, width in column_widths.items():
                ws.column_dimensions[col_letter].width = width
            texto_final = self.context.get("texto_final")
            if texto_final:
                row += 1
                ws.merge_cells(f"A{row}:D{row}")
                cell = ws[f"A{row}"]
                cell.value = self.processar_cabecalho_html(texto_final)
                cell.font = normal_font
                cell.alignment = Alignment(
                    horizontal="left", vertical="top", wrap_text=True
                )
            buffer = BytesIO()
            wb.save(buffer)
            buffer.seek(0)
            for p in temp_image_paths:
                try:
                    if os.path.exists(p):
                        os.unlink(p)
                except Exception:
                    pass
            response = HttpResponse(
                buffer.read(),
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
            response["Content-Disposition"] = (
                f'attachment; filename="{filename}"'
            )
            return response
        except Exception as exc:
            logger.error("Erro ao gerar Excel: %s", exc, exc_info=True)
            raise

    def render_to_docx(
        self,
        cargos_list: Any,
        context: Any,
        texto_final: Any,
        filename: Any = "relatorio_vagas.docx",
    ) -> Any:
        """Gera um arquivo Word (DOCX) mantendo a estrutura hierárquica do.

        Args:
            self: Instância do objeto.
            cargos_list: Lista de cargos com suas DREs e vagas (estrutura.
            context: Contexto com os dados do relatório.
            texto_final: Texto final do relatório.
            filename: Nome do arquivo Word gerado.

        Returns:
            Resultado da operação.

        Raises:
            ImportError: Se ocorrer erro nesta operação.
        """
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx não está instalado. Instale com: pip install python-docx>=1.1.0"  # noqa: E501
            )
        try:
            doc = Document()
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            RGBColor(102, 126, 234)
            RGBColor(52, 73, 94)
            RGBColor(236, 240, 241)
            if context.get("cabecalho_padrao"):
                cabecalho_texto = self.processar_cabecalho_html(
                    context.get("cabecalho_padrao")
                )
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cabecalho_texto)
                run.font.size = Pt(14)
                run.font.bold = True
                doc.add_paragraph()
            if context.get("cabecalho"):
                cabecalho_texto = self.processar_cabecalho_html(
                    context.get("cabecalho")
                )
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cabecalho_texto)
                run.font.size = Pt(14)
                run.font.bold = True
                doc.add_paragraph()
            for cargo in cargos_list:
                cargo_descricao = cargo.get("descricao", "")
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(f"Cargo: {cargo_descricao}")
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                p_pr = p._element.get_or_add_pPr()
                existing_shd = p_pr.find(qn("w:shd"))
                if existing_shd is not None:
                    p_pr.remove(existing_shd)
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(qn("w:fill"), "667eea")
                shading_elm.set(qn("w:val"), "clear")
                p_pr.append(shading_elm)
                for dre in cargo.get("dres", []):
                    dre_nome = dre.get("nome", "")
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(f"DRE - {dre_nome}")
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    p_pr = p._element.get_or_add_pPr()
                    existing_shd = p_pr.find(qn("w:shd"))
                    if existing_shd is not None:
                        p_pr.remove(existing_shd)
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:fill"), "34495e")
                    shading_elm.set(qn("w:val"), "clear")
                    p_pr.append(shading_elm)
                    headers = [
                        "Tipo de unidade",
                        "Unidade",
                        "Vagas Definitivas",
                        "Vagas Precárias",
                    ]
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = "Light Grid Accent 1"
                    header_cells = table.rows[0].cells
                    for i, header in enumerate(headers):
                        cell = header_cells[i]
                        cell.text = header
                        cell.paragraphs[
                            0
                        ].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].runs[0].font.size = Pt(10)
                        tc_pr = cell._element.get_or_add_tcPr()
                        existing_shd = tc_pr.find(qn("w:shd"))
                        if existing_shd is not None:
                            tc_pr.remove(existing_shd)
                        shading_elm = OxmlElement("w:shd")
                        shading_elm.set(qn("w:fill"), "ECF0F1")
                        shading_elm.set(qn("w:val"), "clear")
                        tc_pr.append(shading_elm)
                    for vaga in dre.get("vagas", []):
                        escola = vaga.get("escola", {})
                        row_cells = table.add_row().cells
                        row_cells[0].text = escola.get("tipo_ue", "-")
                        row_cells[1].text = escola.get("nome_oficial", "-")
                        row_cells[2].text = str(
                            vaga.get("vagas_definitivas", 0)
                        )
                        row_cells[3].text = str(vaga.get("vagas_precarias", 0))
                        for i, cell in enumerate(row_cells):
                            if i in [2, 3]:
                                cell.paragraphs[
                                    0
                                ].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                cell.paragraphs[
                                    0
                                ].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            cell.paragraphs[0].runs[0].font.size = Pt(10)
                    doc.add_paragraph()
            if texto_final:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(self.processar_cabecalho_html(texto_final))
                run.font.size = Pt(10)
                doc.add_paragraph()
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            response = HttpResponse(
                buffer.read(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = (
                f'attachment; filename="{filename}"'
            )
            return response
        except Exception as exc:
            logger.error("Erro ao gerar Word: %s", exc, exc_info=True)
            raise
